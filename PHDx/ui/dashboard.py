"""
PHDx Dashboard v3 - PhD Thesis Command Center
Premier Doctoral Writing Tool with Glassmorphism 2.0

Features:
- Focus Canvas (850px centered drafting pane)
- Vertical Stepper navigation with Ring Chart progress
- Logic Glow micro-interactions for Red Thread Engine
- Traffic Light feedback system
"""

import json
import os
import sys
from pathlib import Path

import streamlit as st

# Add core to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from core.citations import ZoteroSentinel, render_sentinel_widget
from core.red_thread import RedThreadEngine
from core.auditor import BrookesAuditor, GoogleDocsPusher
from core.ethics_utils import (
    scrub_text,
    quick_scrub,
    log_ai_usage,
    get_usage_stats,
    get_scrubber
)
from core.supervisor_loop import SupervisorLoop, render_supervisor_notes_widget
from core.feedback_processor import FeedbackProcessor, render_feedback_tab, get_highlight_text
from core.transparency import TransparencyLog, render_declaration_export
from core.secrets_utils import get_secret

# Paths
ROOT_DIR = Path(__file__).parent.parent
DATA_DIR = ROOT_DIR / "data"
DRAFTS_DIR = ROOT_DIR / "drafts"
DNA_PATH = DATA_DIR / "author_dna.json"
CSS_PATH = Path(__file__).parent / "modern_styles.css"

# Page configuration
st.set_page_config(
    page_title="PHDx - PhD Thesis Command Center",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================================================
# LOAD EXTERNAL CSS + GLASSMORPHISM 2.0 ENHANCEMENTS
# ============================================================================
def load_css():
    """Load external CSS and add dynamic styles."""
    # Load external CSS file
    if CSS_PATH.exists():
        with open(CSS_PATH, "r") as f:
            external_css = f.read()
    else:
        external_css = ""

    # Additional dynamic styles for Streamlit components
    dynamic_css = """
    <style>
    /* Load external modern styles */
    """ + external_css + """

    /* ========== PHDX SPECIFIC OVERRIDES ========== */

    /* Glassmorphism 2.0 Background */
    .stApp {
        background: linear-gradient(135deg, #002147 0%, #121212 100%);
        background-attachment: fixed;
    }

    /* Subtle animated gradient overlay */
    .stApp::before {
        content: '';
        position: fixed;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background:
            radial-gradient(ellipse at 20% 20%, rgba(0, 113, 206, 0.08) 0%, transparent 50%),
            radial-gradient(ellipse at 80% 80%, rgba(0, 212, 255, 0.05) 0%, transparent 50%);
        pointer-events: none;
        z-index: 0;
    }

    /* Glass containers */
    .glass-panel, .glass-card, .element-container, [data-testid="stVerticalBlock"] > div {
        backdrop-filter: blur(20px);
        -webkit-backdrop-filter: blur(20px);
    }

    /* Focus Canvas - Central Drafting Pane */
    .focus-canvas-wrapper {
        max-width: 850px;
        margin: 0 auto;
        padding: 2rem;
    }

    .focus-canvas {
        background: rgba(255, 255, 255, 0.03);
        backdrop-filter: blur(20px);
        -webkit-backdrop-filter: blur(20px);
        border: 0.5px solid rgba(255, 255, 255, 0.1);
        border-radius: 20px;
        padding: 2.5rem;
        box-shadow: 0 12px 48px rgba(0, 0, 0, 0.5);
        transition: all 0.4s ease;
    }

    /* Logic Glow - Red Thread Consistency Confirmed */
    .focus-canvas.logic-glow {
        border-color: #00D4FF;
        box-shadow:
            0 12px 48px rgba(0, 0, 0, 0.5),
            0 0 30px rgba(0, 212, 255, 0.3);
        animation: logicPulse 2s ease-in-out;
    }

    @keyframes logicPulse {
        0%, 100% {
            box-shadow: 0 12px 48px rgba(0, 0, 0, 0.5), 0 0 20px rgba(0, 212, 255, 0.2);
        }
        50% {
            box-shadow: 0 12px 48px rgba(0, 0, 0, 0.5), 0 0 50px rgba(0, 212, 255, 0.5);
        }
    }

    /* Prose Typography */
    .prose-content, .stTextArea textarea {
        font-family: 'Source Serif 4', Georgia, 'Times New Roman', serif;
        font-size: 1.1rem;
        line-height: 1.6;
        color: #FAFAFA;
    }

    /* Paragraph with ignored correction */
    .correction-ignored {
        border-bottom: 2px solid #FF5252;
        background: linear-gradient(to bottom, transparent 85%, rgba(255, 82, 82, 0.1) 100%);
        position: relative;
    }

    .correction-ignored::after {
        content: '';
        position: absolute;
        bottom: -4px;
        left: 0;
        right: 0;
        height: 2px;
        background: linear-gradient(90deg, transparent, #FF5252, transparent);
        animation: underlinePulse 2s ease-in-out infinite;
    }

    @keyframes underlinePulse {
        0%, 100% { opacity: 0.5; }
        50% { opacity: 1; }
    }

    /* ========== VERTICAL STEPPER ========== */
    .vertical-stepper {
        display: flex;
        flex-direction: column;
        gap: 0.5rem;
        padding: 1rem;
    }

    .stepper-item {
        display: flex;
        align-items: center;
        gap: 1rem;
        padding: 1rem 1.25rem;
        background: rgba(255, 255, 255, 0.03);
        backdrop-filter: blur(20px);
        -webkit-backdrop-filter: blur(20px);
        border: 0.5px solid rgba(255, 255, 255, 0.1);
        border-radius: 12px;
        cursor: pointer;
        transition: all 0.25s ease;
    }

    .stepper-item:hover {
        background: rgba(255, 255, 255, 0.06);
        border-color: #0071ce;
        transform: translateX(4px);
    }

    .stepper-item.active {
        background: rgba(0, 113, 206, 0.15);
        border-color: #0071ce;
        box-shadow: 0 0 20px rgba(0, 113, 206, 0.2);
    }

    .stepper-item.completed {
        border-left: 3px solid #00C853;
    }

    .stepper-number {
        width: 28px;
        height: 28px;
        border-radius: 50%;
        background: rgba(255, 255, 255, 0.05);
        border: 1px solid rgba(255, 255, 255, 0.2);
        display: flex;
        align-items: center;
        justify-content: center;
        font-family: 'Inter', sans-serif;
        font-weight: 600;
        font-size: 0.8rem;
        flex-shrink: 0;
    }

    .stepper-item.active .stepper-number {
        background: #0071ce;
        border-color: #0071ce;
    }

    .stepper-item.completed .stepper-number {
        background: #00C853;
        border-color: #00C853;
    }

    .stepper-label {
        font-family: 'Inter', sans-serif;
        font-size: 0.85rem;
        font-weight: 500;
        color: #FAFAFA;
    }

    .stepper-progress {
        font-size: 0.7rem;
        color: #6B7280;
        margin-top: 2px;
    }

    /* ========== RING CHART ========== */
    .ring-chart-container {
        position: relative;
        width: 140px;
        height: 140px;
        margin: 1.5rem auto;
    }

    .ring-chart {
        width: 100%;
        height: 100%;
        transform: rotate(-90deg);
    }

    .ring-chart-bg {
        fill: none;
        stroke: rgba(255, 255, 255, 0.1);
        stroke-width: 10;
    }

    .ring-chart-progress {
        fill: none;
        stroke: #0071ce;
        stroke-width: 10;
        stroke-linecap: round;
        stroke-dasharray: 377;
        stroke-dashoffset: 377;
        transition: stroke-dashoffset 1s ease-out;
        filter: drop-shadow(0 0 8px rgba(0, 113, 206, 0.5));
    }

    .ring-chart-progress.milestone-25 { stroke: #FFB300; }
    .ring-chart-progress.milestone-50 { stroke: #00D4FF; }
    .ring-chart-progress.milestone-75 { stroke: #00C853; }
    .ring-chart-progress.milestone-100 {
        stroke: #00C853;
        filter: drop-shadow(0 0 12px rgba(0, 200, 83, 0.7));
    }

    .ring-chart-center {
        position: absolute;
        top: 50%;
        left: 50%;
        transform: translate(-50%, -50%);
        text-align: center;
    }

    .ring-chart-value {
        font-family: 'Inter', sans-serif;
        font-size: 1.3rem;
        font-weight: 700;
        color: #FAFAFA;
    }

    .ring-chart-label {
        font-family: 'Inter', sans-serif;
        font-size: 0.6rem;
        color: #6B7280;
        text-transform: uppercase;
        letter-spacing: 0.1em;
    }

    /* ========== SIDEBAR STYLING ========== */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, rgba(0, 33, 71, 0.98) 0%, rgba(18, 18, 18, 0.99) 100%);
        backdrop-filter: blur(20px);
        -webkit-backdrop-filter: blur(20px);
        border-right: 0.5px solid rgba(255, 255, 255, 0.1);
    }

    [data-testid="stSidebar"] > div:first-child {
        padding-top: 1rem;
    }

    /* PHDx Logo */
    .phdx-logo {
        font-family: 'Inter', sans-serif;
        font-size: 2rem;
        font-weight: 700;
        background: linear-gradient(135deg, #0071ce 0%, #00D4FF 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        text-align: center;
        margin-bottom: 0.5rem;
    }

    .phdx-tagline {
        font-family: 'Inter', sans-serif;
        font-size: 0.7rem;
        color: #6B7280;
        text-align: center;
        text-transform: uppercase;
        letter-spacing: 0.15em;
        margin-bottom: 1.5rem;
    }

    /* ========== BUTTONS ========== */
    .stButton > button {
        font-family: 'Inter', sans-serif;
        font-weight: 500;
        background: rgba(255, 255, 255, 0.03);
        backdrop-filter: blur(20px);
        -webkit-backdrop-filter: blur(20px);
        border: 0.5px solid rgba(255, 255, 255, 0.1);
        border-radius: 10px;
        color: #FAFAFA;
        transition: all 0.25s ease;
    }

    .stButton > button:hover {
        background: rgba(255, 255, 255, 0.06);
        border-color: #0071ce;
        box-shadow: 0 0 20px rgba(0, 113, 206, 0.2);
        transform: translateY(-1px);
    }

    /* Primary button */
    .stButton > button[kind="primary"] {
        background: linear-gradient(135deg, #0071ce 0%, #0058a3 100%);
        border: none;
        box-shadow: 0 4px 15px rgba(0, 113, 206, 0.3);
    }

    .stButton > button[kind="primary"]:hover {
        background: linear-gradient(135deg, #0082e6 0%, #0071ce 100%);
        box-shadow: 0 6px 25px rgba(0, 113, 206, 0.4);
    }

    /* ========== TEXT INPUTS ========== */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea {
        font-family: 'Source Serif 4', Georgia, serif;
        font-size: 1.1rem;
        line-height: 1.6;
        background: rgba(255, 255, 255, 0.03);
        backdrop-filter: blur(20px);
        -webkit-backdrop-filter: blur(20px);
        border: 0.5px solid rgba(255, 255, 255, 0.1);
        border-radius: 12px;
        color: #FAFAFA;
        padding: 1rem;
    }

    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus {
        border-color: #0071ce;
        box-shadow: 0 0 0 3px rgba(0, 113, 206, 0.15);
    }

    /* ========== METRICS ========== */
    [data-testid="stMetricValue"] {
        font-family: 'Inter', sans-serif;
        font-weight: 700;
        color: #0071ce !important;
    }

    [data-testid="stMetricLabel"] {
        font-family: 'Inter', sans-serif;
        color: #6B7280 !important;
        text-transform: uppercase;
        letter-spacing: 0.05em;
        font-size: 0.75rem;
    }

    /* ========== TABS ========== */
    .stTabs [data-baseweb="tab-list"] {
        background: rgba(255, 255, 255, 0.02);
        backdrop-filter: blur(20px);
        -webkit-backdrop-filter: blur(20px);
        border-radius: 12px;
        border: 0.5px solid rgba(255, 255, 255, 0.1);
        padding: 4px;
        gap: 4px;
    }

    .stTabs [data-baseweb="tab"] {
        font-family: 'Inter', sans-serif;
        color: #6B7280;
        border-radius: 10px;
        font-weight: 500;
    }

    .stTabs [aria-selected="true"] {
        background: rgba(0, 113, 206, 0.15);
        color: #0071ce !important;
        border: 1px solid rgba(0, 113, 206, 0.3);
    }

    /* ========== GLASS CARDS ========== */
    .glass-card {
        background: rgba(255, 255, 255, 0.03);
        backdrop-filter: blur(20px);
        -webkit-backdrop-filter: blur(20px);
        border-radius: 12px;
        border: 0.5px solid rgba(255, 255, 255, 0.1);
        padding: 1.5rem;
        margin-bottom: 1rem;
        box-shadow: 0 8px 32px 0 rgba(0, 0, 0, 0.4);
    }

    /* ========== TEXT COLORS ========== */
    .stMarkdown, .stText, p, span, label {
        color: #e0e0e0 !important;
    }

    h1, h2, h3, h4, h5, h6 {
        color: #ffffff !important;
        font-family: 'Inter', sans-serif;
    }

    /* ========== MAIN HEADER ========== */
    .main-header {
        font-size: 2.5rem;
        font-weight: 700;
        background: linear-gradient(135deg, #0071ce 0%, #00D4FF 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        margin-bottom: 0.25rem;
        letter-spacing: -0.02em;
        text-align: center;
    }

    .sub-header {
        font-size: 0.9rem;
        color: #6B7280 !important;
        margin-bottom: 2rem;
        font-weight: 400;
        text-align: center;
    }

    /* ========== STATUS BAR ========== */
    .status-bar {
        display: flex;
        gap: 1rem;
        margin-bottom: 1.5rem;
        padding: 0.5rem;
    }

    .status-item {
        background: rgba(255, 255, 255, 0.03);
        backdrop-filter: blur(20px);
        -webkit-backdrop-filter: blur(20px);
        border-radius: 12px;
        border: 0.5px solid rgba(255, 255, 255, 0.1);
        padding: 1rem 1.5rem;
        flex: 1;
        text-align: center;
        transition: all 0.3s ease;
    }

    .status-item:hover {
        border-color: rgba(0, 113, 206, 0.5);
        box-shadow: 0 0 20px rgba(0, 113, 206, 0.15);
    }

    .status-label {
        font-size: 0.7rem;
        color: #6B7280;
        text-transform: uppercase;
        letter-spacing: 0.1em;
        margin-bottom: 0.25rem;
        font-family: 'Inter', sans-serif;
    }

    .status-value {
        font-size: 1.4rem;
        font-weight: 700;
        color: #0071ce;
        font-family: 'Inter', sans-serif;
    }

    .status-value.success { color: #00C853; }
    .status-value.warning { color: #FFB300; }

    /* ========== SCROLLBAR ========== */
    ::-webkit-scrollbar {
        width: 8px;
        height: 8px;
    }

    ::-webkit-scrollbar-track {
        background: rgba(255, 255, 255, 0.02);
    }

    ::-webkit-scrollbar-thumb {
        background: rgba(0, 113, 206, 0.3);
        border-radius: 4px;
    }

    ::-webkit-scrollbar-thumb:hover {
        background: rgba(0, 113, 206, 0.5);
    }

    /* ========== PROGRESS BARS ========== */
    .stProgress > div > div > div {
        background: linear-gradient(90deg, #0071ce 0%, #00D4FF 100%);
        border-radius: 6px;
    }

    /* ========== TRAFFIC LIGHT BADGES ========== */
    .traffic-red {
        background: rgba(255, 82, 82, 0.15);
        border: 1px solid rgba(255, 82, 82, 0.3);
        color: #FF5252;
        padding: 0.25rem 0.75rem;
        border-radius: 20px;
        font-size: 0.8rem;
    }

    .traffic-amber {
        background: rgba(255, 179, 0, 0.15);
        border: 1px solid rgba(255, 179, 0, 0.3);
        color: #FFB300;
        padding: 0.25rem 0.75rem;
        border-radius: 20px;
        font-size: 0.8rem;
    }

    .traffic-green {
        background: rgba(0, 200, 83, 0.15);
        border: 1px solid rgba(0, 200, 83, 0.3);
        color: #00C853;
        padding: 0.25rem 0.75rem;
        border-radius: 20px;
        font-size: 0.8rem;
    }

    /* ========== ETHICS BADGE ========== */
    .ethics-badge {
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
        background: linear-gradient(135deg, rgba(0, 200, 83, 0.15) 0%, rgba(0, 150, 60, 0.15) 100%);
        backdrop-filter: blur(20px);
        -webkit-backdrop-filter: blur(20px);
        border: 1px solid rgba(0, 200, 83, 0.3);
        border-radius: 20px;
        padding: 0.4rem 1rem;
        font-size: 0.85rem;
        color: #00C853;
        font-weight: 500;
        font-family: 'Inter', sans-serif;
    }

    /* ========== COMPLEXITY GAUGE ========== */
    .complexity-gauge {
        background: rgba(255, 255, 255, 0.03);
        backdrop-filter: blur(20px);
        -webkit-backdrop-filter: blur(20px);
        border-radius: 12px;
        border: 0.5px solid rgba(255, 255, 255, 0.1);
        padding: 1rem;
        margin-top: 1rem;
    }

    .complexity-bar-container {
        background: rgba(255, 255, 255, 0.05);
        border-radius: 8px;
        height: 12px;
        position: relative;
        overflow: hidden;
    }

    .complexity-bar {
        height: 100%;
        border-radius: 8px;
        transition: width 0.5s ease;
    }

    .complexity-bar.optimal { background: linear-gradient(90deg, #00C853, #00E676); }
    .complexity-bar.acceptable { background: linear-gradient(90deg, #0071ce, #00D4FF); }
    .complexity-bar.warning { background: linear-gradient(90deg, #FF9800, #FFB300); }
    .complexity-bar.danger { background: linear-gradient(90deg, #D32F2F, #FF5252); }

    /* ========== THEMATIC HEATMAP ========== */
    .heatmap-container {
        background: rgba(255, 255, 255, 0.03);
        backdrop-filter: blur(20px);
        -webkit-backdrop-filter: blur(20px);
        border-radius: 12px;
        border: 0.5px solid rgba(255, 255, 255, 0.1);
        padding: 1.5rem;
        margin: 1rem 0;
    }

    .heatmap-bar {
        display: flex;
        align-items: center;
        gap: 1rem;
        margin-bottom: 0.75rem;
    }

    .heatmap-label {
        min-width: 120px;
        font-family: 'Inter', sans-serif;
        font-size: 0.85rem;
        color: #B0B0B0;
    }

    .heatmap-track {
        flex: 1;
        height: 24px;
        background: rgba(255, 255, 255, 0.05);
        border-radius: 6px;
        overflow: hidden;
    }

    .heatmap-fill {
        height: 100%;
        border-radius: 6px;
        transition: width 0.5s ease;
    }

    .word-count {
        min-width: 100px;
        text-align: right;
        font-family: 'Inter', sans-serif;
        font-size: 0.8rem;
        color: #6B7280;
    }
    </style>
    """

    st.markdown(dynamic_css, unsafe_allow_html=True)

# Load CSS at the start
load_css()


# ============================================================================
# SESSION STATE INITIALIZATION
# ============================================================================
if "drafting_text" not in st.session_state:
    st.session_state.drafting_text = ""
if "current_chapter" not in st.session_state:
    st.session_state.current_chapter = "Chapter 1: Introduction"
if "zotero_sentinel" not in st.session_state:
    st.session_state.zotero_sentinel = ZoteroSentinel()
if "red_thread_engine" not in st.session_state:
    st.session_state.red_thread_engine = RedThreadEngine()
if "brookes_auditor" not in st.session_state:
    st.session_state.brookes_auditor = BrookesAuditor()
if "google_docs_pusher" not in st.session_state:
    st.session_state.google_docs_pusher = GoogleDocsPusher()
if "audit_report" not in st.session_state:
    st.session_state.audit_report = None
if "ethically_scanned" not in st.session_state:
    st.session_state.ethically_scanned = False
if "last_scrub_report" not in st.session_state:
    st.session_state.last_scrub_report = None
if "supervisor_loop" not in st.session_state:
    st.session_state.supervisor_loop = SupervisorLoop()
if "feedback_processor" not in st.session_state:
    st.session_state.feedback_processor = FeedbackProcessor()
if "highlight_text" not in st.session_state:
    st.session_state.highlight_text = None
if "transparency_log" not in st.session_state:
    st.session_state.transparency_log = TransparencyLog()
if "logic_glow" not in st.session_state:
    st.session_state.logic_glow = False
if "ignored_corrections" not in st.session_state:
    st.session_state.ignored_corrections = []


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================
def load_author_dna() -> dict | None:
    """Load the author DNA profile if it exists."""
    if DNA_PATH.exists():
        with open(DNA_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return None


def calculate_flesch_kincaid(text: str) -> dict:
    """
    Calculate Flesch-Kincaid Grade Level for text complexity analysis.

    Formula: 0.39 * (words/sentences) + 11.8 * (syllables/words) - 15.59

    PhD target: Grade 14-16 (doctoral level)
    Warning: Below 10 = too informal

    Returns:
        dict: {
            "grade_level": float,
            "word_count": int,
            "sentence_count": int,
            "syllable_count": int,
            "avg_words_per_sentence": float,
            "avg_syllables_per_word": float,
            "status": "optimal" | "acceptable" | "warning" | "danger",
            "message": str
        }
    """
    import re

    if not text or len(text.strip()) < 50:
        return {
            "grade_level": 0,
            "word_count": 0,
            "sentence_count": 0,
            "syllable_count": 0,
            "avg_words_per_sentence": 0,
            "avg_syllables_per_word": 0,
            "status": "warning",
            "message": "Not enough text to analyze"
        }

    def count_syllables(word: str) -> int:
        """Count syllables in a word using vowel patterns."""
        word = word.lower().strip()
        if not word:
            return 0

        # Handle common exceptions
        if len(word) <= 3:
            return 1

        # Count vowel groups
        vowels = "aeiouy"
        count = 0
        prev_vowel = False

        for char in word:
            is_vowel = char in vowels
            if is_vowel and not prev_vowel:
                count += 1
            prev_vowel = is_vowel

        # Adjust for silent e
        if word.endswith('e') and count > 1:
            count -= 1

        # Adjust for -le endings
        if word.endswith('le') and len(word) > 2 and word[-3] not in vowels:
            count += 1

        # Ensure at least 1 syllable
        return max(1, count)

    # Clean text
    clean_text = re.sub(r'[^\w\s.!?]', '', text)

    # Count sentences (split on .!?)
    sentences = re.split(r'[.!?]+', clean_text)
    sentences = [s.strip() for s in sentences if s.strip() and len(s.split()) > 2]
    sentence_count = max(1, len(sentences))

    # Count words
    words = re.findall(r'\b[a-zA-Z]+\b', clean_text)
    word_count = len(words)

    if word_count == 0:
        return {
            "grade_level": 0,
            "word_count": 0,
            "sentence_count": 0,
            "syllable_count": 0,
            "avg_words_per_sentence": 0,
            "avg_syllables_per_word": 0,
            "status": "warning",
            "message": "No valid words found"
        }

    # Count syllables
    syllable_count = sum(count_syllables(word) for word in words)

    # Calculate averages
    avg_words_per_sentence = word_count / sentence_count
    avg_syllables_per_word = syllable_count / word_count

    # Flesch-Kincaid Grade Level formula
    grade_level = (0.39 * avg_words_per_sentence) + (11.8 * avg_syllables_per_word) - 15.59
    grade_level = round(max(0, grade_level), 1)

    # Determine status and message for PhD context
    if 14 <= grade_level <= 16:
        status = "optimal"
        message = "Doctoral level complexity - optimal for PhD"
    elif 12 <= grade_level < 14:
        status = "acceptable"
        message = "Graduate level - acceptable, could be more complex"
    elif 16 < grade_level <= 18:
        status = "acceptable"
        message = "Highly complex - consider readability for examiners"
    elif grade_level <= 10:
        status = "danger"
        message = "Tone too informal for Doctoral level"
    elif grade_level < 12:
        status = "warning"
        message = "Below doctoral standard - increase complexity"
    else:
        status = "warning"
        message = "Very high complexity - may hinder comprehension"

    return {
        "grade_level": grade_level,
        "word_count": word_count,
        "sentence_count": sentence_count,
        "syllable_count": syllable_count,
        "avg_words_per_sentence": round(avg_words_per_sentence, 1),
        "avg_syllables_per_word": round(avg_syllables_per_word, 2),
        "status": status,
        "message": message
    }


def render_complexity_gauge(text: str):
    """Render the ComplexityGauge widget for Flesch-Kincaid analysis."""
    analysis = calculate_flesch_kincaid(text)

    grade = analysis["grade_level"]
    status = analysis["status"]

    # Calculate bar width (scale: 0-20 grade level)
    bar_width = min(100, max(0, (grade / 20) * 100))

    # Target zone: 14-16 on 0-20 scale = 70-80%
    target_left = 70
    target_width = 10

    st.markdown(f"""
    <div class="complexity-gauge">
        <div class="complexity-gauge-header">
            <span class="complexity-gauge-title">Flesch-Kincaid Grade Level</span>
            <span class="complexity-gauge-value {status}">{grade}</span>
        </div>
        <div class="complexity-bar-container">
            <div class="complexity-target-zone" style="left: {target_left}%; width: {target_width}%;"></div>
            <div class="complexity-bar {status}" style="width: {bar_width}%;"></div>
        </div>
        <div class="complexity-labels">
            <span>0 (Simple)</span>
            <span>10</span>
            <span>14-16 (PhD)</span>
            <span>20 (Complex)</span>
        </div>
        {"<div class='complexity-warning'>⚠️ " + analysis['message'] + "</div>" if status in ['warning', 'danger'] else ""}
        {"<div class='complexity-optimal'>✓ " + analysis['message'] + "</div>" if status == 'optimal' else ""}
    </div>
    """, unsafe_allow_html=True)

    return analysis


def inject_sheets_data() -> dict | None:
    """
    Pull active row from Google Sheets and generate synthesis.

    Returns synthesized text in author's DNA voice.
    """
    import anthropic
    from dotenv import load_dotenv

    load_dotenv()

    # Check for Google Sheets credentials
    sheets_url = os.getenv("GOOGLE_SHEETS_URL")
    if not sheets_url:
        return {"error": "GOOGLE_SHEETS_URL not configured in .env"}

    # Load author DNA for voice matching
    dna = load_author_dna()
    if not dna:
        return {"error": "Author DNA profile not found. Run dna_engine.py first."}

    # Get Claude client
    api_key = get_secret("ANTHROPIC_API_KEY")
    if not api_key:
        return {"error": "ANTHROPIC_API_KEY not configured"}

    try:
        import gspread
        from google.oauth2.service_account import Credentials

        # Try to connect to Google Sheets
        creds_path = os.getenv("GOOGLE_CREDENTIALS_PATH", "credentials.json")

        if not Path(creds_path).exists():
            return {"error": f"Google credentials not found at {creds_path}"}

        scopes = [
            "https://www.googleapis.com/auth/spreadsheets.readonly",
            "https://www.googleapis.com/auth/drive.readonly"
        ]

        creds = Credentials.from_service_account_file(creds_path, scopes=scopes)
        gc = gspread.authorize(creds)

        # Open sheet and get active/first row with data
        sheet = gc.open_by_url(sheets_url).sheet1
        records = sheet.get_all_records()

        if not records:
            return {"error": "No data found in sheet"}

        # Get the first row (or could be configured to get "active" row)
        active_row = records[0]

        # Prepare DNA context
        dna_context = json.dumps(dna.get("claude_deep_analysis", {}), indent=2)

        # Generate synthesis with Claude
        client = anthropic.Anthropic(api_key=api_key)

        # Ethics scrubbing: Anonymize data before sending to AI
        data_to_synthesize = json.dumps(active_row, indent=2)
        scrub_result = scrub_text(data_to_synthesize)
        scrubbed_data = scrub_result["scrubbed_text"]
        was_scrubbed = scrub_result["total_redactions"] > 0

        prompt = f"""You are writing a critical synthesis for a PhD thesis. You must match the author's unique writing style.

AUTHOR'S WRITING DNA:
{dna_context}

Linguistic fingerprint:
- Average sentence length: {dna.get('sentence_complexity', {}).get('average_length', 20)} words
- Hedging density: {dna.get('hedging_analysis', {}).get('hedging_density_per_1000_words', 5)} per 1000 words
- Preferred transitions: {', '.join(dna.get('transition_vocabulary', {}).get('preferred_categories', ['contrast', 'addition'])[:3])}

DATA TO SYNTHESIZE:
{scrubbed_data}

Write a 300-word critical synthesis of this data that:
1. Matches the author's sentence structure and complexity
2. Uses their characteristic hedging language
3. Employs their preferred transition vocabulary
4. Maintains their academic voice and tone

Write ONLY the synthesis paragraph, no meta-commentary."""

        # Log AI usage
        log_ai_usage(
            action_type="generate_synthesis",
            data_source="google_sheets",
            prompt=prompt,
            was_scrubbed=was_scrubbed,
            redactions_count=scrub_result["total_redactions"]
        )

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1024,
            messages=[{"role": "user", "content": prompt}]
        )

        return {
            "success": True,
            "data_row": active_row,
            "synthesis": response.content[0].text,
            "ethically_scanned": True,
            "redactions": scrub_result["total_redactions"]
        }

    except ImportError:
        return {"error": "gspread not installed. Run: pip install gspread"}
    except Exception as e:
        return {"error": str(e)}


# ============================================================================
# SIDEBAR - Vertical Stepper with Ring Chart
# ============================================================================
def get_total_word_count() -> int:
    """Calculate total word count from all draft documents."""
    total = 0
    if DRAFTS_DIR.exists():
        from docx import Document
        for docx_file in DRAFTS_DIR.glob("*.docx"):
            try:
                doc = Document(docx_file)
                text = " ".join([p.text for p in doc.paragraphs])
                total += len(text.split())
            except Exception:
                pass
    return total


def render_ring_chart(current_words: int, target_words: int = 80000) -> str:
    """Generate SVG Ring Chart for thesis progress."""
    percentage = min(100, (current_words / target_words) * 100)
    # SVG circle: circumference = 2 * pi * r = 2 * 3.14159 * 60 = 377
    circumference = 377
    offset = circumference - (percentage / 100) * circumference

    # Determine milestone class
    if percentage >= 100:
        milestone_class = "milestone-100"
    elif percentage >= 75:
        milestone_class = "milestone-75"
    elif percentage >= 50:
        milestone_class = "milestone-50"
    elif percentage >= 25:
        milestone_class = "milestone-25"
    else:
        milestone_class = ""

    return f"""
    <div class="ring-chart-container">
        <svg class="ring-chart" viewBox="0 0 140 140">
            <circle class="ring-chart-bg" cx="70" cy="70" r="60"/>
            <circle class="ring-chart-progress {milestone_class}"
                    cx="70" cy="70" r="60"
                    style="stroke-dashoffset: {offset}"/>
        </svg>
        <div class="ring-chart-center">
            <div class="ring-chart-value">{percentage:.0f}%</div>
            <div class="ring-chart-label">of 80k</div>
        </div>
    </div>
    """


def render_vertical_stepper(active_section: str = "drafting") -> str:
    """Generate Vertical Stepper navigation HTML."""
    sections = [
        {"id": "drafting", "label": "Drafting", "icon": "✍️", "progress": "Active"},
        {"id": "feedback", "label": "Feedback", "icon": "📝", "progress": "3 items"},
        {"id": "dna", "label": "DNA Profile", "icon": "🧬", "progress": "Generated"},
        {"id": "chapters", "label": "Chapters", "icon": "📚", "progress": "6 chapters"},
        {"id": "progress", "label": "Progress", "icon": "📈", "progress": "On track"},
    ]

    html = '<div class="vertical-stepper">'
    for i, section in enumerate(sections, 1):
        is_active = section["id"] == active_section
        is_completed = i < sections.index(next(s for s in sections if s["id"] == active_section)) + 1 if active_section else False

        state_class = "active" if is_active else ("completed" if is_completed else "")

        html += f"""
        <div class="stepper-item {state_class}">
            <div class="stepper-number">{section["icon"]}</div>
            <div>
                <div class="stepper-label">{section["label"]}</div>
                <div class="stepper-progress">{section["progress"]}</div>
            </div>
        </div>
        """
    html += '</div>'
    return html


def render_sidebar():
    """Render the sidebar with Vertical Stepper navigation and Ring Chart."""
    with st.sidebar:
        # PHDx Logo and tagline
        st.markdown('<div class="phdx-logo">PHDx</div>', unsafe_allow_html=True)
        st.markdown('<div class="phdx-tagline">Doctoral Command Center</div>', unsafe_allow_html=True)

        # Ring Chart - 80,000 word progress
        total_words = get_total_word_count()
        st.markdown(render_ring_chart(total_words, 80000), unsafe_allow_html=True)

        # Word count display
        st.markdown(f"""
        <div style="text-align: center; margin-bottom: 1.5rem;">
            <span style="font-family: 'Inter', sans-serif; font-size: 1.5rem; font-weight: 700; color: #0071ce;">
                {total_words:,}
            </span>
            <span style="font-family: 'Inter', sans-serif; font-size: 0.9rem; color: #6B7280;">
                / 80,000 words
            </span>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("---")

        # Vertical Stepper Navigation
        st.markdown(render_vertical_stepper("drafting"), unsafe_allow_html=True)

        st.markdown("---")

        # Red Thread Engine Status
        st.markdown("#### 🔴 Red Thread Engine")
        engine = st.session_state.red_thread_engine
        stats = engine.get_stats()

        col1, col2 = st.columns(2)
        with col1:
            st.metric("Indexed", stats["total_paragraphs"])
        with col2:
            backend = "Cloud" if stats.get("backend") == "pinecone" else "Local"
            st.metric("Backend", backend)

        if st.button("🔄 Re-index", use_container_width=True):
            with st.spinner("Indexing..."):
                result = engine.index_drafts_folder()
                st.success(f"Indexed {result['paragraphs_indexed']} paragraphs")
                st.session_state.logic_glow = True

        st.markdown("---")

        # Zotero Sentinel (Compact)
        sentinel = st.session_state.zotero_sentinel
        status = "✅ Connected" if sentinel.connected else ("🧪 Mock Mode" if sentinel.mock_mode else "❌ Offline")
        st.markdown(f"#### 📚 Zotero: {status}")

        if sentinel.connected or sentinel.mock_mode:
            st.metric("Sources", len(sentinel.items_cache) if sentinel.items_cache else 5)

        st.markdown("---")

        # Settings (Compact)
        with st.expander("⚙️ Settings"):
            st.selectbox(
                "Citation Style",
                ["Harvard (Cite Them Right)", "APA 7th", "Chicago", "MLA"],
                index=0,
                key="citation_style"
            )
            st.selectbox(
                "Language",
                ["British English", "American English"],
                index=0,
                key="language"
            )

        # DNA Profile status
        dna_exists = DNA_PATH.exists()
        if dna_exists:
            st.markdown('<span class="ethics-badge">🧬 DNA Active</span>', unsafe_allow_html=True)
        else:
            if st.button("🧬 Generate DNA Profile", use_container_width=True):
                st.info("Run: python core/dna_engine.py")


# ============================================================================
# MAIN CONTENT - Focus Canvas Layout
# ============================================================================
def render_drafting_pane():
    """Render the main drafting pane with Focus Canvas (850px max-width)."""

    # Centered header
    st.markdown('<p class="main-header">PHDx</p>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">Premier Doctoral Writing Tool | Oxford Brookes Standards</p>', unsafe_allow_html=True)

    # Status bar with glassmorphic cards
    dna_profile = load_author_dna()
    total_words = get_total_word_count()

    st.markdown(f"""
    <div class="status-bar">
        <div class="status-item">
            <div class="status-label">Words Written</div>
            <div class="status-value">{total_words:,}</div>
        </div>
        <div class="status-item">
            <div class="status-label">Documents</div>
            <div class="status-value">{len(dna_profile['metadata']['documents_analyzed']) if dna_profile else '--'}</div>
        </div>
        <div class="status-item">
            <div class="status-label">Avg Sentence</div>
            <div class="status-value">{dna_profile['sentence_complexity']['average_length'] if dna_profile else '--'}</div>
        </div>
        <div class="status-item">
            <div class="status-label">Progress</div>
            <div class="status-value success">{min(100, int(total_words / 800))}%</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Main tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "✍️ Drafting",
        "📝 Supervisor Feedback",
        "🧬 DNA Profile",
        "📚 Chapters",
        "📈 Progress"
    ])

    with tab1:
        render_drafting_tab()

    with tab2:
        render_feedback_tab(st.session_state.feedback_processor)

    with tab3:
        render_dna_tab(dna_profile)

    with tab4:
        render_chapters_tab()

    with tab5:
        render_progress_tab()


def render_drafting_tab():
    """Render the main drafting interface with Focus Canvas."""

    # Determine if Logic Glow is active (Red Thread consistency confirmed)
    logic_glow_class = "logic-glow" if st.session_state.logic_glow else ""

    # Start Focus Canvas wrapper
    st.markdown(f'<div class="focus-canvas-wrapper"><div class="focus-canvas {logic_glow_class}">', unsafe_allow_html=True)

    # ========== SUPERVISOR FEEDBACK HIGHLIGHT ==========
    highlight_text = get_highlight_text()
    if highlight_text:
        st.markdown(f"""
        <div style="background: linear-gradient(135deg, rgba(255, 152, 0, 0.15) 0%, rgba(255, 193, 7, 0.15) 100%);
                    border: 1px solid rgba(255, 152, 0, 0.4);
                    border-left: 4px solid #ff9800;
                    border-radius: 8px; padding: 1rem; margin-bottom: 1rem;">
            <div style="display: flex; justify-content: space-between; align-items: center;">
                <div>
                    <strong style="color: #ffc107;">🎯 Feedback Target:</strong>
                    <span style="color: rgba(224, 224, 224, 0.9);"> Look for text containing:</span>
                </div>
            </div>
            <div style="background: rgba(0, 0, 0, 0.2); border-radius: 6px; padding: 0.75rem; margin-top: 0.5rem;
                        font-family: monospace; color: #ffc107;">
                "{highlight_text[:100]}{'...' if len(highlight_text) > 100 else ''}"
            </div>
            <div style="margin-top: 0.5rem; font-size: 0.85rem; color: rgba(224, 224, 224, 0.7);">
                Use Ctrl+F in the text area below to find this passage, or check the Supervisor Feedback tab.
            </div>
        </div>
        """, unsafe_allow_html=True)

        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("Clear Highlight", key="clear_hl_drafting"):
                st.session_state.highlight_text = None
                st.rerun()

    # ========== LOGIC GLOW STATUS ==========
    if st.session_state.logic_glow:
        st.markdown("""
        <div style="background: rgba(0, 212, 255, 0.1); border: 1px solid rgba(0, 212, 255, 0.3);
                    border-radius: 8px; padding: 0.75rem; margin-bottom: 1rem; display: flex;
                    align-items: center; gap: 0.5rem;">
            <span style="color: #00D4FF; font-size: 1.2rem;">✓</span>
            <span style="color: #00D4FF; font-weight: 500;">Red Thread: Logical consistency confirmed</span>
        </div>
        """, unsafe_allow_html=True)

    # ========== ETHICALLY SCANNED BADGE ==========
    if st.session_state.ethically_scanned:
        st.markdown("""
        <div style="display: flex; align-items: center; margin-bottom: 1rem;">
            <span class="ethics-badge">
                <span class="ethics-badge-icon">🛡️</span>
                Ethically Scanned
            </span>
        </div>
        """, unsafe_allow_html=True)

        # Show scrub stats if available
        if st.session_state.last_scrub_report:
            report = st.session_state.last_scrub_report
            st.markdown(f"""
            <div class="ethics-stats">
                PII items redacted: <strong>{report.get('total_redactions', 0)}</strong> |
                Data anonymized before AI processing
            </div>
            """, unsafe_allow_html=True)

    st.markdown("### Drafting Pane")

    # Chapter selector
    col1, col2 = st.columns([2, 1])

    with col1:
        chapter = st.selectbox(
            "Working on:",
            [
                "Chapter 1: Introduction",
                "Chapter 2: Literature Review",
                "Chapter 3: Methodology",
                "Chapter 4: Results",
                "Chapter 5: Discussion",
                "Chapter 6: Conclusion",
                "Abstract",
                "Free Writing"
            ],
            key="chapter_select"
        )
        st.session_state.current_chapter = chapter

    with col2:
        st.selectbox(
            "AI Mode:",
            ["Style Match", "Expand", "Summarize", "Critique", "Off"]
        )

    # Writing area
    text_input = st.text_area(
        "Start writing...",
        height=400,
        placeholder="Begin drafting your thesis here. Your writing will be analyzed against your DNA profile for style consistency.",
        label_visibility="collapsed",
        key="draft_text"
    )
    st.session_state.drafting_text = text_input

    # ========== COMPLEXITY GAUGE ==========
    # Show Flesch-Kincaid analysis for text complexity
    if text_input and len(text_input.strip()) >= 50:
        render_complexity_gauge(text_input)
    elif text_input and len(text_input.strip()) > 0:
        st.markdown("""
        <div class="complexity-gauge" style="opacity: 0.6;">
            <div class="complexity-gauge-header">
                <span class="complexity-gauge-title">Flesch-Kincaid Grade Level</span>
                <span class="complexity-gauge-value" style="color: rgba(224, 224, 224, 0.5);">--</span>
            </div>
            <div class="complexity-bar-container">
                <div class="complexity-target-zone" style="left: 70%; width: 10%;"></div>
            </div>
            <div class="complexity-labels">
                <span>0 (Simple)</span>
                <span>10</span>
                <span>14-16 (PhD)</span>
                <span>20 (Complex)</span>
            </div>
            <div class="complexity-warning">⚠️ Write at least 50 characters for complexity analysis</div>
        </div>
        """, unsafe_allow_html=True)

    # Action buttons row 1
    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.button("🔍 Check Style", use_container_width=True)

    with col2:
        st.button("✨ AI Assist", use_container_width=True)

    with col3:
        st.button("💾 Save Draft", use_container_width=True)

    with col4:
        st.button("📤 Export", use_container_width=True)

    # Action buttons row 2 - Thematic Squeeze
    st.markdown("---")
    st.markdown("#### 📊 Thematic Squeeze")

    col1, col2 = st.columns([1, 2])

    with col1:
        if st.button("📥 Inject Data from Sheets", use_container_width=True, type="primary"):
            with st.spinner("Fetching data and generating synthesis..."):
                result = inject_sheets_data()

                if result.get("error"):
                    st.error(result["error"])
                else:
                    st.session_state.sheets_synthesis = result
                    # Set ethically scanned state
                    if result.get("ethically_scanned"):
                        st.session_state.ethically_scanned = True
                        st.session_state.last_scrub_report = {"total_redactions": result.get("redactions", 0)}

    with col2:
        if "sheets_synthesis" in st.session_state and st.session_state.sheets_synthesis.get("success"):
            st.success("Synthesis generated in your Authorial DNA voice!")

    # Display synthesis result
    if "sheets_synthesis" in st.session_state and st.session_state.sheets_synthesis.get("success"):
        result = st.session_state.sheets_synthesis

        with st.expander("📋 Source Data Row", expanded=False):
            st.json(result["data_row"])

        st.markdown("#### Generated Synthesis (300 words, DNA-matched)")
        st.markdown(f'<div class="glass-card">{result["synthesis"]}</div>', unsafe_allow_html=True)

        if st.button("📋 Copy to Draft", use_container_width=True):
            st.session_state.drafting_text = result["synthesis"]
            st.rerun()

    # Red Thread continuity check
    st.markdown("---")
    st.markdown("#### 🔴 Continuity Check")

    if text_input and len(text_input) > 100:
        if st.button("Check for Contradictions", use_container_width=True):
            with st.spinner("Analyzing against your thesis corpus..."):
                engine = st.session_state.red_thread_engine

                # Use enhanced UI report with thematic threshold warning
                ui_report = engine.get_consistency_report_for_ui(text_input)

                # Thematic Threshold Warning - display if similarity < 50%
                if ui_report.get("low_consistency_warning"):
                    warning = ui_report["low_consistency_warning"]
                    st.markdown(f"""
                    <div style="background: linear-gradient(135deg, rgba(255, 152, 0, 0.2) 0%, rgba(255, 87, 34, 0.2) 100%);
                                border: 1px solid rgba(255, 152, 0, 0.5);
                                border-left: 4px solid #ff9800;
                                border-radius: 8px; padding: 1rem; margin-bottom: 1rem;">
                        <div style="display: flex; align-items: center; gap: 0.5rem; margin-bottom: 0.5rem;">
                            <span style="font-size: 1.5rem;">⚠️</span>
                            <strong style="color: #ffc107; font-size: 1.1rem;">{warning['message']}</strong>
                        </div>
                        <div style="color: rgba(224, 224, 224, 0.9); font-size: 0.9rem;">
                            Average similarity: <strong>{warning['avg_similarity']:.1f}%</strong> (below 50% threshold)
                        </div>
                        <div style="color: rgba(224, 224, 224, 0.7); font-size: 0.85rem; margin-top: 0.5rem;">
                            {warning['recommendation']}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)

                # Show overall score
                score = ui_report.get("score", 0)
                score_label = ui_report.get("score_label", "Unknown")
                score_color = "#00c853" if score >= 80 else "#ffc107" if score >= 50 else "#f44336"

                st.markdown(f"**Consistency Score:** <span style='color: {score_color};'>{score}/100</span> ({score_label})", unsafe_allow_html=True)

                # Show issues
                if ui_report.get("issue_count", 0) == 0:
                    st.success("✓ No contradictions detected")
                else:
                    for severity in ["high", "medium", "low"]:
                        issues = ui_report["issues_by_severity"].get(severity, [])
                        for issue in issues:
                            severity_icons = {"high": "🔴", "medium": "🟡", "low": "🟢"}
                            icon = severity_icons.get(severity, "⚪")
                            st.warning(f"{icon} **{issue.get('type', 'Issue').title()}**: {issue.get('description', '')}")
                            if issue.get("recommendation"):
                                st.info(f"💡 {issue['recommendation']}")

                # Show summary
                if ui_report.get("summary"):
                    st.caption(ui_report["summary"])
    else:
        st.info("Write at least 100 characters to check for contradictions")

    # ========== CRITIQUE MODE - OXFORD BROOKES AUDIT ==========
    st.markdown("---")
    st.markdown("#### 🎓 Critique Mode - Oxford Brookes Audit")
    st.markdown("*Evaluate your draft against PhD marking criteria: Originality, Criticality, Rigour*")

    col1, col2 = st.columns([1, 2])

    with col1:
        if st.button("🔍 Audit Draft", use_container_width=True, type="primary",
                     disabled=len(text_input) < 100):
            with st.spinner("Evaluating against Oxford Brookes criteria..."):
                # Ethics scrubbing before AI processing
                scrub_result = scrub_text(text_input)
                scrubbed_text = scrub_result["scrubbed_text"]
                was_scrubbed = scrub_result["total_redactions"] > 0

                # Build the audit prompt for logging
                audit_prompt = f"Audit draft for chapter: {chapter}. Text length: {len(text_input)} chars."

                # Log AI usage
                log_ai_usage(
                    action_type="audit_draft",
                    data_source="drafting_pane",
                    prompt=audit_prompt,
                    was_scrubbed=was_scrubbed,
                    redactions_count=scrub_result["total_redactions"]
                )

                auditor = st.session_state.brookes_auditor
                report = auditor.audit_draft(scrubbed_text, chapter)
                st.session_state.audit_report = report

                # Set ethically scanned state
                st.session_state.ethically_scanned = True
                st.session_state.last_scrub_report = scrub_result

    with col2:
        if st.session_state.audit_report and st.session_state.audit_report.get("status") == "success":
            grade = st.session_state.audit_report["overall_grade"]
            level_emoji = {
                "excellent": "🌟",
                "good": "✅",
                "satisfactory": "📊",
                "needs_improvement": "⚠️",
                "unsatisfactory": "❌"
            }
            emoji = level_emoji.get(grade["level"], "📋")
            st.success(f"{emoji} Last Audit: **{grade['score']}/100** ({grade['level'].replace('_', ' ').title()})")

    # Display audit report if available
    if st.session_state.audit_report:
        report = st.session_state.audit_report

        if report.get("error"):
            st.error(f"Audit Error: {report['error']}")
        elif report.get("status") == "success":
            # Overall grade display
            grade = report["overall_grade"]
            level_colors = {
                "excellent": "#00c853",
                "good": "#0071ce",
                "satisfactory": "#ffc107",
                "needs_improvement": "#ff9800",
                "unsatisfactory": "#f44336"
            }
            color = level_colors.get(grade["level"], "#e0e0e0")

            st.markdown(f"""
            <div class="glass-card" style="border-left: 4px solid {color};">
                <h4 style="margin-bottom: 0.5rem;">Overall Grade: <span style="color: {color};">{grade['score']}/100</span></h4>
                <p style="color: rgba(224, 224, 224, 0.8); margin-bottom: 0.5rem;">{grade['level'].replace('_', ' ').title()}</p>
                <p style="font-size: 0.9rem;">{grade['descriptor']}</p>
            </div>
            """, unsafe_allow_html=True)

            # Criteria breakdown
            with st.expander("📊 Criteria Breakdown", expanded=True):
                scores = report["criteria_scores"]

                col1, col2, col3 = st.columns(3)

                with col1:
                    st.markdown("**Originality (35%)**")
                    st.metric("Score", f"{scores['originality']['score']}/100")
                    st.caption(scores['originality']['feedback'])

                with col2:
                    st.markdown("**Critical Analysis (35%)**")
                    st.metric("Score", f"{scores['criticality']['score']}/100")
                    st.caption(scores['criticality']['feedback'])

                with col3:
                    st.markdown("**Rigour (30%)**")
                    st.metric("Score", f"{scores['rigour']['score']}/100")
                    st.caption(scores['rigour']['feedback'])

            # Strengths and improvements
            with st.expander("💪 Strengths & Areas for Improvement"):
                col1, col2 = st.columns(2)

                with col1:
                    st.markdown("**Strengths:**")
                    for s in report.get("strengths", []):
                        st.markdown(f"✓ {s}")

                with col2:
                    st.markdown("**Areas for Improvement:**")
                    for a in report.get("areas_for_improvement", []):
                        st.markdown(f"→ {a}")

            # Recommendations
            with st.expander("💡 Recommendations"):
                for i, rec in enumerate(report.get("specific_recommendations", []), 1):
                    st.markdown(f"{i}. {rec}")

            # Examiner summary
            with st.expander("📝 Examiner Summary"):
                st.markdown(report.get("examiner_summary", ""))

    # ========== PUSH TO GOOGLE DOC ==========
    st.markdown("---")
    st.markdown("#### 📄 Push to Google Doc")

    # Google Doc ID input
    doc_id = st.text_input(
        "Google Doc ID",
        value=os.getenv("GOOGLE_DOC_ID", ""),
        placeholder="Enter the Document ID from Google Docs URL",
        help="Find this in your Google Doc URL: docs.google.com/document/d/{DOC_ID}/edit"
    )

    section_title = st.text_input(
        "Section Title (optional)",
        placeholder="e.g., Chapter 3 Draft - Methodology"
    )

    col1, col2 = st.columns([1, 2])

    with col1:
        push_disabled = not doc_id or len(text_input) < 50
        if st.button("📤 Push to Google Doc", use_container_width=True, type="primary",
                     disabled=push_disabled):
            with st.spinner("Pushing to Google Doc with PHDx-Verified timestamp..."):
                pusher = st.session_state.google_docs_pusher
                result = pusher.push_to_doc(
                    doc_id=doc_id,
                    text=text_input,
                    section_title=section_title,
                    include_timestamp=True
                )

                if result["success"]:
                    st.success(f"✅ Successfully pushed {result['characters_added']:,} characters!")
                    st.markdown(f"[Open Document]({result['doc_url']})")
                else:
                    st.error(f"Failed: {result.get('error', 'Unknown error')}")

    with col2:
        if not doc_id:
            st.info("Enter a Google Doc ID to enable push")
        elif len(text_input) < 50:
            st.info("Write at least 50 characters to push")
        else:
            st.markdown(f"Ready to push **{len(text_input):,}** characters with PHDx-Verified timestamp")

    # Close Focus Canvas wrapper
    st.markdown('</div></div>', unsafe_allow_html=True)

    # Reset logic glow after render (one-time animation)
    if st.session_state.logic_glow:
        st.session_state.logic_glow = False


def render_dna_tab(dna_profile: dict | None):
    """Render the DNA profile analysis view."""
    st.markdown("### Your Writing DNA")

    if not dna_profile:
        st.warning("No DNA profile found. Add .docx files to the `/drafts` folder and run the DNA engine.")
        st.code("python core/dna_engine.py", language="bash")
        return

    # Metrics overview
    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("#### Sentence Complexity")
        st.metric(
            "Average Length",
            f"{dna_profile['sentence_complexity']['average_length']} words"
        )

        dist = dna_profile['sentence_complexity']['length_distribution']
        if dist:
            st.bar_chart(dist)

    with col2:
        st.markdown("#### Hedging Usage")
        st.metric(
            "Density",
            f"{dna_profile['hedging_analysis']['hedging_density_per_1000_words']}/1000 words"
        )

        top_hedges = dict(list(dna_profile['hedging_analysis']['phrases_found'].items())[:5])
        if top_hedges:
            st.markdown("**Top phrases:**")
            for phrase, count in top_hedges.items():
                st.markdown(f"- *{phrase}*: {count}")

    with col3:
        st.markdown("#### Transitions")
        st.metric(
            "Density",
            f"{dna_profile['transition_vocabulary']['transition_density_per_1000_words']}/1000 words"
        )

        prefs = dna_profile['transition_vocabulary']['preferred_categories']
        if prefs:
            st.markdown("**Preferred types:**")
            for cat in prefs:
                st.markdown(f"- {cat.replace('_', ' ').title()}")

    # Claude analysis
    st.markdown("---")
    st.markdown("#### Deep Analysis")

    claude_analysis = dna_profile.get('claude_deep_analysis', {})
    if "error" in claude_analysis:
        st.warning(f"Claude analysis unavailable: {claude_analysis['error']}")
    elif "raw_analysis" in claude_analysis:
        st.markdown(claude_analysis['raw_analysis'])
    else:
        st.json(claude_analysis)


def render_chapters_tab():
    """Render the chapters overview."""
    st.markdown("### Thesis Chapters")

    chapters = [
        {"name": "Introduction", "status": "In Progress", "words": 0, "target": 8000},
        {"name": "Literature Review", "status": "Not Started", "words": 0, "target": 20000},
        {"name": "Methodology", "status": "Not Started", "words": 0, "target": 15000},
        {"name": "Results", "status": "Not Started", "words": 0, "target": 15000},
        {"name": "Discussion", "status": "Not Started", "words": 0, "target": 15000},
        {"name": "Conclusion", "status": "Not Started", "words": 0, "target": 7000},
    ]

    for i, chapter in enumerate(chapters, 1):
        with st.expander(f"Chapter {i}: {chapter['name']} ({chapter['status']})"):
            progress = chapter['words'] / chapter['target'] if chapter['target'] > 0 else 0
            st.progress(progress)
            st.markdown(f"**Words:** {chapter['words']:,} / {chapter['target']:,}")

            col1, col2 = st.columns(2)
            with col1:
                st.button(f"📝 Edit", key=f"edit_{i}", use_container_width=True)
            with col2:
                st.button(f"📊 Analyze", key=f"analyze_{i}", use_container_width=True)


def render_progress_tab():
    """Render progress tracking with Thematic Heatmap."""
    st.markdown("### Thesis Progress")

    # Chapter word count targets (typical PhD thesis structure)
    chapters = [
        {"name": "Introduction", "short": "Intro", "target": 8000, "color": "#4fc3f7"},
        {"name": "Literature Review", "short": "Lit Review", "target": 20000, "color": "#81c784"},
        {"name": "Methodology", "short": "Methods", "target": 15000, "color": "#ffb74d"},
        {"name": "Results", "short": "Results", "target": 15000, "color": "#ba68c8"},
        {"name": "Discussion", "short": "Discussion", "target": 15000, "color": "#f06292"},
        {"name": "Conclusion", "short": "Conclusion", "target": 7000, "color": "#4db6ac"},
    ]

    # Calculate actual word counts from draft files
    chapter_keywords = {
        "Introduction": ["intro", "introduction"],
        "Literature Review": ["lit", "literature", "review"],
        "Methodology": ["method", "methodology"],
        "Results": ["result", "findings", "data"],
        "Discussion": ["discussion", "analysis"],
        "Conclusion": ["conclusion", "concluding"],
    }

    # Scan drafts folder for word counts
    for chapter in chapters:
        chapter["words"] = 0
        keywords = chapter_keywords.get(chapter["name"], [])

        if DRAFTS_DIR.exists():
            for docx_file in DRAFTS_DIR.glob("*.docx"):
                filename_lower = docx_file.stem.lower()
                if any(kw in filename_lower for kw in keywords):
                    try:
                        from docx import Document
                        doc = Document(docx_file)
                        text = " ".join([p.text for p in doc.paragraphs])
                        chapter["words"] += len(text.split())
                    except Exception:
                        pass

    # Calculate totals
    total_words = sum(c["words"] for c in chapters)
    target_words = 80000

    # Overall progress
    progress = total_words / target_words if target_words > 0 else 0
    st.markdown(f"**Overall: {total_words:,} / {target_words:,} words ({progress*100:.1f}%)**")
    st.progress(min(progress, 1.0))

    # ========== THEMATIC HEATMAP ==========
    st.markdown("---")
    st.markdown("#### Thematic Heatmap")
    st.caption("Chapter progress towards 80,000-word goal")

    # Create horizontal bar chart using HTML/CSS for glassmorphic style
    heatmap_html = """
    <style>
    .heatmap-container {
        background: rgba(30, 30, 30, 0.6);
        border-radius: 12px;
        padding: 1.5rem;
        backdrop-filter: blur(10px);
    }
    .chapter-row {
        display: flex;
        align-items: center;
        margin-bottom: 0.75rem;
    }
    .chapter-label {
        width: 100px;
        font-size: 0.85rem;
        color: rgba(255, 255, 255, 0.9);
        text-align: right;
        padding-right: 1rem;
    }
    .bar-container {
        flex: 1;
        height: 28px;
        background: rgba(255, 255, 255, 0.1);
        border-radius: 6px;
        overflow: hidden;
        position: relative;
    }
    .bar-fill {
        height: 100%;
        border-radius: 6px;
        transition: width 0.5s ease;
        position: relative;
    }
    .bar-text {
        position: absolute;
        right: 8px;
        top: 50%;
        transform: translateY(-50%);
        font-size: 0.75rem;
        color: rgba(255, 255, 255, 0.9);
        font-weight: 500;
        text-shadow: 0 1px 2px rgba(0,0,0,0.5);
    }
    .target-marker {
        position: absolute;
        right: 0;
        top: 0;
        height: 100%;
        width: 2px;
        background: rgba(255, 255, 255, 0.3);
    }
    .word-count {
        width: 120px;
        text-align: right;
        font-size: 0.8rem;
        color: rgba(255, 255, 255, 0.7);
        padding-left: 0.75rem;
    }
    </style>
    <div class="heatmap-container">
    """

    for chapter in chapters:
        percent = (chapter["words"] / chapter["target"] * 100) if chapter["target"] > 0 else 0
        bar_width = min(percent, 100)  # Cap at 100% visually

        # Color intensity based on progress
        opacity = 0.4 + (min(percent, 100) / 100) * 0.6

        heatmap_html += f"""
        <div class="chapter-row">
            <div class="chapter-label">{chapter["short"]}</div>
            <div class="bar-container">
                <div class="bar-fill" style="width: {bar_width}%; background: linear-gradient(90deg, {chapter["color"]}88, {chapter["color"]});">
                    <span class="bar-text">{percent:.0f}%</span>
                </div>
                <div class="target-marker"></div>
            </div>
            <div class="word-count">{chapter["words"]:,} / {chapter["target"]:,}</div>
        </div>
        """

    heatmap_html += "</div>"

    st.markdown(heatmap_html, unsafe_allow_html=True)

    # Legend
    st.caption("Bar shows progress toward chapter target. Right edge = 100% of chapter goal.")

    # Milestones
    st.markdown("---")
    st.markdown("#### Milestones")

    milestones = [
        {"name": "Literature Review Complete", "done": False},
        {"name": "Methodology Approved", "done": False},
        {"name": "Data Collection Complete", "done": False},
        {"name": "First Draft Complete", "done": False},
        {"name": "Supervisor Review", "done": False},
        {"name": "Final Submission", "done": False},
    ]

    for milestone in milestones:
        icon = "✅" if milestone['done'] else "⬜"
        st.markdown(f"{icon} {milestone['name']}")

    # AI Transparency Declaration Export
    st.markdown("---")
    render_declaration_export(st.session_state.transparency_log)


# ============================================================================
# MAIN ENTRY POINT
# ============================================================================
def main():
    """Main application entry point."""
    render_sidebar()
    render_drafting_pane()


if __name__ == "__main__":
    main()
