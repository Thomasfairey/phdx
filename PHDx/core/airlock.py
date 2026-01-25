"""
Airlock - Data Connector Module

Google OAuth 2.0 authentication and document loading utilities for the PhD
writing assistant. Provides secure access to Google Drive, Docs, and Sheets,
as well as local file parsing capabilities.
"""

from pathlib import Path
from typing import Optional

from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError


# OAuth 2.0 Scopes
SCOPES = [
    'https://www.googleapis.com/auth/drive.file',
    'https://www.googleapis.com/auth/documents.readonly',
    'https://www.googleapis.com/auth/spreadsheets.readonly',
]

# Configuration paths
CONFIG_DIR = Path(__file__).parent.parent / 'config'
CLIENT_SECRET_PATH = CONFIG_DIR / 'client_secret.json'
TOKEN_PATH = CONFIG_DIR / 'token.json'


def authenticate_user() -> Credentials:
    """
    Authenticate user via Google OAuth 2.0.

    Loads credentials from config/token.json if valid, otherwise launches
    the browser-based OAuth flow to obtain new credentials.

    Returns:
        google.oauth2.credentials.Credentials: Authenticated credentials object.

    Raises:
        FileNotFoundError: If client_secret.json is not found.
    """
    creds = None

    # Ensure config directory exists
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)

    # Check for existing token
    if TOKEN_PATH.exists():
        creds = Credentials.from_authorized_user_file(str(TOKEN_PATH), SCOPES)

    # If no valid credentials, initiate OAuth flow
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            # Refresh expired credentials
            creds.refresh(Request())
        else:
            # Run full OAuth flow
            if not CLIENT_SECRET_PATH.exists():
                raise FileNotFoundError(
                    f"Client secret not found at {CLIENT_SECRET_PATH}. "
                    "Please download OAuth credentials from Google Cloud Console."
                )

            flow = InstalledAppFlow.from_client_secrets_file(
                str(CLIENT_SECRET_PATH), SCOPES
            )
            creds = flow.run_local_server(port=0)

        # Save credentials for future use
        with open(TOKEN_PATH, 'w') as token_file:
            token_file.write(creds.to_json())

    return creds


def list_recent_docs(limit: int = 10) -> list[dict]:
    """
    List the most recently modified Google Docs and Sheets.

    Args:
        limit: Maximum number of documents to return (default 10).

    Returns:
        List of dictionaries with keys: 'name', 'id', 'type'.
    """
    creds = authenticate_user()
    service = build('drive', 'v3', credentials=creds)

    # Query for Google Docs and Sheets, ordered by modified time
    query = (
        "mimeType='application/vnd.google-apps.document' or "
        "mimeType='application/vnd.google-apps.spreadsheet'"
    )

    try:
        results = service.files().list(
            q=query,
            pageSize=limit,
            orderBy='modifiedTime desc',
            fields='files(id, name, mimeType)'
        ).execute()

        files = results.get('files', [])

        # Map MIME types to friendly names
        mime_type_map = {
            'application/vnd.google-apps.document': 'doc',
            'application/vnd.google-apps.spreadsheet': 'sheet',
        }

        return [
            {
                'name': f['name'],
                'id': f['id'],
                'type': mime_type_map.get(f['mimeType'], 'unknown'),
            }
            for f in files
        ]

    except HttpError as error:
        print(f"An error occurred: {error}")
        return []


def load_google_doc(doc_id: str) -> str:
    """
    Load the full text content of a Google Doc.

    Args:
        doc_id: The Google Doc ID.

    Returns:
        Extracted text content as a string.
    """
    creds = authenticate_user()
    service = build('docs', 'v1', credentials=creds)

    try:
        document = service.documents().get(documentId=doc_id).execute()

        # Parse document structure to extract text
        text_content = _extract_text_from_doc(document)
        return text_content

    except HttpError as error:
        print(f"An error occurred: {error}")
        return ""


def _extract_text_from_doc(document: dict) -> str:
    """
    Extract plain text from a Google Docs JSON structure.

    Args:
        document: The document JSON from the Docs API.

    Returns:
        Concatenated text content.
    """
    text_parts = []

    body = document.get('body', {})
    content = body.get('content', [])

    for element in content:
        if 'paragraph' in element:
            paragraph = element['paragraph']
            elements = paragraph.get('elements', [])

            for elem in elements:
                if 'textRun' in elem:
                    text = elem['textRun'].get('content', '')
                    text_parts.append(text)

        elif 'table' in element:
            # Handle table content
            table = element['table']
            for row in table.get('tableRows', []):
                for cell in row.get('tableCells', []):
                    cell_content = cell.get('content', [])
                    for cell_elem in cell_content:
                        if 'paragraph' in cell_elem:
                            for p_elem in cell_elem['paragraph'].get('elements', []):
                                if 'textRun' in p_elem:
                                    text = p_elem['textRun'].get('content', '')
                                    text_parts.append(text)

    return ''.join(text_parts)


def load_local_file(uploaded_file) -> str:
    """
    Load text content from a Streamlit UploadedFile object.

    Supports .docx (Word) and .pdf files.

    Args:
        uploaded_file: Streamlit UploadedFile object.

    Returns:
        Extracted text content as a string.

    Raises:
        ValueError: If file type is not supported.
    """
    filename = uploaded_file.name.lower()

    if filename.endswith('.docx'):
        return _load_docx(uploaded_file)
    elif filename.endswith('.pdf'):
        return _load_pdf(uploaded_file)
    elif filename.endswith('.txt'):
        return uploaded_file.read().decode('utf-8')
    else:
        raise ValueError(
            f"Unsupported file type: {filename}. "
            "Supported formats: .docx, .pdf, .txt"
        )


def _load_docx(uploaded_file) -> str:
    """
    Extract text from a .docx file.

    Args:
        uploaded_file: File-like object containing .docx data.

    Returns:
        Extracted text content.
    """
    from docx import Document
    from io import BytesIO

    # Reset file pointer if needed
    uploaded_file.seek(0)

    doc = Document(BytesIO(uploaded_file.read()))

    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return '\n\n'.join(paragraphs)


def _load_pdf(uploaded_file) -> str:
    """
    Extract text from a .pdf file.

    Args:
        uploaded_file: File-like object containing PDF data.

    Returns:
        Extracted text content.
    """
    from pypdf import PdfReader
    from io import BytesIO

    # Reset file pointer if needed
    uploaded_file.seek(0)

    reader = PdfReader(BytesIO(uploaded_file.read()))

    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)

    return '\n\n'.join(text_parts)


# Utility functions for Streamlit integration
def get_auth_status() -> dict:
    """
    Check the current authentication status.

    Returns:
        Dictionary with 'authenticated' bool and 'message' string.
    """
    if not CLIENT_SECRET_PATH.exists():
        return {
            'authenticated': False,
            'message': 'OAuth not configured. Please add client_secret.json to config/'
        }

    if not TOKEN_PATH.exists():
        return {
            'authenticated': False,
            'message': 'Not authenticated. Click to sign in with Google.'
        }

    try:
        creds = Credentials.from_authorized_user_file(str(TOKEN_PATH), SCOPES)
        if creds.valid:
            return {
                'authenticated': True,
                'message': 'Connected to Google Drive'
            }
        elif creds.expired and creds.refresh_token:
            return {
                'authenticated': False,
                'message': 'Token expired. Click to refresh.'
            }
        else:
            return {
                'authenticated': False,
                'message': 'Invalid credentials. Please re-authenticate.'
            }
    except Exception as e:
        return {
            'authenticated': False,
            'message': f'Error checking credentials: {str(e)}'
        }


def get_credentials() -> Optional[Credentials]:
    """
    Get current credentials if valid, otherwise return None.

    Returns:
        Credentials object if authenticated and valid, None otherwise.
    """
    if not TOKEN_PATH.exists():
        return None

    try:
        creds = Credentials.from_authorized_user_file(str(TOKEN_PATH), SCOPES)
        if creds.valid:
            return creds
        elif creds.expired and creds.refresh_token:
            creds.refresh(Request())
            # Save refreshed credentials
            with open(TOKEN_PATH, 'w') as token_file:
                token_file.write(creds.to_json())
            return creds
    except Exception:
        pass

    return None


def clear_credentials() -> None:
    """Remove stored OAuth token to force re-authentication."""
    if TOKEN_PATH.exists():
        TOKEN_PATH.unlink()


# =============================================================================
# API COMPATIBILITY FUNCTIONS
# =============================================================================

def get_document_text(doc_id: str) -> dict:
    """
    Get document text content with error handling for API use.

    Args:
        doc_id: The Google Doc ID.

    Returns:
        Dictionary with 'success', 'text', and optional 'error' keys.
    """
    try:
        text = load_google_doc(doc_id)
        if text:
            return {"success": True, "text": text}
        else:
            return {"success": False, "text": "", "error": "Document is empty or could not be read"}
    except FileNotFoundError as e:
        return {"success": False, "text": "", "error": str(e)}
    except HttpError as e:
        return {"success": False, "text": "", "error": f"Google API error: {e}"}
    except Exception as e:
        return {"success": False, "text": "", "error": str(e)}


def get_user_info() -> dict:
    """
    Get current user authentication info for API use.

    Returns:
        Dictionary with 'email', 'name', 'authenticated', and optional 'mock' keys.
    """
    status = get_auth_status()

    if not status['authenticated']:
        return {
            "email": "",
            "name": "",
            "authenticated": False,
            "mock": None
        }

    # If authenticated, try to get user info from the API
    try:
        creds = get_credentials()
        if creds:
            # Build People API to get user info
            try:
                from googleapiclient.discovery import build
                service = build('people', 'v1', credentials=creds)
                profile = service.people().get(
                    resourceName='people/me',
                    personFields='names,emailAddresses'
                ).execute()

                names = profile.get('names', [{}])
                emails = profile.get('emailAddresses', [{}])

                return {
                    "email": emails[0].get('value', '') if emails else '',
                    "name": names[0].get('displayName', '') if names else '',
                    "authenticated": True,
                    "mock": False
                }
            except Exception:
                # Fall back to basic authenticated response
                return {
                    "email": "authenticated@user",
                    "name": "Authenticated User",
                    "authenticated": True,
                    "mock": False
                }
    except Exception:
        pass

    return {
        "email": "",
        "name": "",
        "authenticated": False,
        "mock": None
    }


def update_google_doc(doc_id: str, content: str, section_title: Optional[str] = None) -> dict:
    """
    Append content to a Google Doc.

    Args:
        doc_id: The Google Doc ID.
        content: Text content to append.
        section_title: Optional section heading to add before content.

    Returns:
        Dictionary with 'success', 'doc_url', and optional 'error' keys.
    """
    from datetime import datetime

    try:
        creds = get_credentials()
        if not creds:
            return {
                "success": False,
                "doc_url": None,
                "error": "Not authenticated. Please authenticate first."
            }

        service = build('docs', 'v1', credentials=creds)

        # Get document to find end index
        doc = service.documents().get(documentId=doc_id).execute()
        body = doc.get('body', {})
        doc_content = body.get('content', [])

        if not doc_content:
            end_index = 1
        else:
            end_index = doc_content[-1].get('endIndex', 1) - 1

        # Build content to insert
        content_parts = []
        content_parts.append("\n\n---\n\n")

        # Add timestamp
        timestamp_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        content_parts.append(f"[PHDx-Verified: {timestamp_str}]\n\n")

        # Add section title if provided
        if section_title:
            content_parts.append(f"## {section_title}\n\n")

        content_parts.append(content)
        full_content = "".join(content_parts)

        # Create insert request
        requests = [
            {
                'insertText': {
                    'location': {'index': end_index},
                    'text': full_content
                }
            }
        ]

        service.documents().batchUpdate(
            documentId=doc_id,
            body={'requests': requests}
        ).execute()

        return {
            "success": True,
            "doc_url": f"https://docs.google.com/document/d/{doc_id}/edit"
        }

    except HttpError as e:
        return {
            "success": False,
            "doc_url": None,
            "error": f"Google API error: {e}"
        }
    except Exception as e:
        return {
            "success": False,
            "doc_url": None,
            "error": str(e)
        }
