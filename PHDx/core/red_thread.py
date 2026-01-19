"""
Red Thread Engine - Hierarchical Logic Check (Map-Reduce Architecture)

Implements a three-level argument continuity analysis for PhD theses:
    Level 1 (Map): Extract key claims from each chunk using TIER_SPEED
    Level 2 (Reduce): Synthesize chapter abstracts from claim extractions
    Level 3 (Audit): Verify argument continuity using TIER_LOGIC

This replaces simple vector similarity with true "Argument Continuity" checking.
"""

import hashlib
import json
import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from dotenv import load_dotenv
from pydantic import BaseModel, Field

# Import the Triad Strategy LLM Gateway
from core.llm_gateway import (
    get_router,
    EngineType,
    ModelTier,
    LLMResponse,
)
from core.vector_store import get_vector_store, VectorStoreBase
from core.secrets_utils import get_secret

# Set up logging
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

# Paths
ROOT_DIR = Path(__file__).parent.parent
DRAFTS_DIR = ROOT_DIR / "drafts"
DATA_DIR = ROOT_DIR / "data"
CHROMA_DIR = DATA_DIR / "chroma_db"

# Collection names
COLLECTION_NAME = "thesis_paragraphs"
CHAPTERS_COLLECTION = "thesis_chapters"

# Chunking configuration
CHUNK_SIZE_TOKENS = 2000
CHUNK_OVERLAP_TOKENS = 200
CHARS_PER_TOKEN = 4  # Rough estimate


# =============================================================================
# Pydantic Response Models (Frontend Compatible)
# =============================================================================

class ChunkClaim(BaseModel):
    """Key claims extracted from a single chunk."""
    chunk_index: int = Field(..., description="Index of the source chunk")
    key_claims: list[str] = Field(default_factory=list, description="List of key claims")
    methodology_notes: Optional[str] = Field(None, description="Any methodology mentioned")
    evidence_cited: list[str] = Field(default_factory=list, description="Evidence or citations")


class ChapterAbstract(BaseModel):
    """Synthesized abstract for a chapter."""
    chapter_id: str = Field(..., description="Chapter identifier (e.g., 'ch1', 'ch5')")
    chapter_title: str = Field(..., description="Chapter title or description")
    core_argument: str = Field(..., description="The main argument of this chapter")
    key_claims: list[str] = Field(default_factory=list, description="Consolidated key claims")
    research_question: Optional[str] = Field(None, description="Research question if Ch1")
    conclusion_answer: Optional[str] = Field(None, description="Conclusion answer if Ch5")


class MissingLink(BaseModel):
    """A gap in the argument chain."""
    from_chapter: str = Field(..., description="Source chapter")
    to_chapter: str = Field(..., description="Target chapter")
    description: str = Field(..., description="What's missing")
    severity: str = Field("medium", description="high/medium/low")
    suggestion: str = Field(..., description="How to address this gap")


class GraphNode(BaseModel):
    """Node for visual argument graph."""
    id: str = Field(..., description="Node identifier")
    label: str = Field(..., description="Display label")
    type: str = Field(..., description="Type: 'question', 'argument', 'evidence', 'conclusion'")
    chapter: str = Field(..., description="Source chapter")


class GraphEdge(BaseModel):
    """Edge for visual argument graph."""
    source: str = Field(..., description="Source node ID")
    target: str = Field(..., description="Target node ID")
    label: Optional[str] = Field(None, description="Edge label")
    strength: float = Field(1.0, description="Connection strength 0-1")


class RedThreadResult(BaseModel):
    """
    Complete Red Thread analysis result.

    Frontend compatible with RedThreadModule.tsx
    """
    # Core result fields (frontend compatible)
    continuity_score: int = Field(..., description="0-100 continuity score")
    thread_status: str = Field(..., description="'solid' or 'broken'")
    status: str = Field(..., description="Human readable status")
    analysis: str = Field(..., description="Detailed analysis text")

    # Detailed findings
    missing_links: list[MissingLink] = Field(default_factory=list)
    chapter_abstracts: list[ChapterAbstract] = Field(default_factory=list)

    # Visual graph data
    visual_graph_nodes: list[GraphNode] = Field(default_factory=list)
    visual_graph_edges: list[GraphEdge] = Field(default_factory=list)

    # Metadata
    processing_metadata: dict = Field(default_factory=dict)

    class Config:
        extra = "allow"


# =============================================================================
# Chunking Utilities
# =============================================================================

def estimate_tokens(text: str) -> int:
    """Estimate token count from text."""
    return len(text) // CHARS_PER_TOKEN


def chunk_text_sliding_window(
    text: str,
    chunk_size: int = CHUNK_SIZE_TOKENS,
    overlap: int = CHUNK_OVERLAP_TOKENS,
) -> list[dict]:
    """
    Split text into overlapping chunks using a sliding window.

    Args:
        text: The full text to chunk
        chunk_size: Target size in tokens
        overlap: Overlap between chunks in tokens

    Returns:
        List of chunk dicts with 'text', 'start_char', 'end_char', 'token_estimate'
    """
    chunk_chars = chunk_size * CHARS_PER_TOKEN
    overlap_chars = overlap * CHARS_PER_TOKEN

    chunks = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = min(start + chunk_chars, text_len)

        # Try to break at sentence boundary
        if end < text_len:
            # Look for sentence end within last 20% of chunk
            search_start = start + int(chunk_chars * 0.8)
            sentence_end = text.rfind('. ', search_start, end)
            if sentence_end > search_start:
                end = sentence_end + 1

        chunk_text = text[start:end].strip()

        if chunk_text:
            chunks.append({
                'index': len(chunks),
                'text': chunk_text,
                'start_char': start,
                'end_char': end,
                'token_estimate': estimate_tokens(chunk_text),
            })

        # Move window with overlap
        start = end - overlap_chars
        if start >= text_len - overlap_chars:
            break

    return chunks


def identify_chapter_boundaries(text: str) -> list[dict]:
    """
    Attempt to identify chapter boundaries in text.

    Looks for patterns like "Chapter 1", "CHAPTER ONE", "1. Introduction", etc.

    Returns:
        List of chapter dicts with 'chapter_id', 'title', 'start_char', 'end_char'
    """
    # Common chapter patterns
    patterns = [
        r'(?i)^(chapter\s+(\d+|one|two|three|four|five|six|seven|eight|nine|ten)[:\s]*(.*))',
        r'(?i)^((\d+)\.\s+([A-Z][^.\n]+))',
        r'(?i)^(part\s+(\d+|one|two|three|four|five)[:\s]*(.*))',
    ]

    chapters = []

    for pattern in patterns:
        for match in re.finditer(pattern, text, re.MULTILINE):
            full_match = match.group(1)
            chapter_num = match.group(2)
            title = match.group(3).strip() if match.group(3) else ""

            # Normalize chapter number
            word_to_num = {
                'one': '1', 'two': '2', 'three': '3', 'four': '4', 'five': '5',
                'six': '6', 'seven': '7', 'eight': '8', 'nine': '9', 'ten': '10'
            }
            if chapter_num.lower() in word_to_num:
                chapter_num = word_to_num[chapter_num.lower()]

            chapters.append({
                'chapter_id': f"ch{chapter_num}",
                'chapter_num': int(chapter_num) if chapter_num.isdigit() else 0,
                'title': title or f"Chapter {chapter_num}",
                'start_char': match.start(),
                'match_text': full_match,
            })

    # Sort by position and dedupe
    chapters = sorted(chapters, key=lambda x: x['start_char'])

    # Set end positions
    for i, ch in enumerate(chapters):
        if i + 1 < len(chapters):
            ch['end_char'] = chapters[i + 1]['start_char']
        else:
            ch['end_char'] = len(text)

    return chapters


# =============================================================================
# Hierarchical Logic Check Engine
# =============================================================================

class HierarchicalLogicChecker:
    """
    Implements the three-level Map-Reduce argument continuity check.

    Level 1: Extract key claims from chunks (TIER_SPEED)
    Level 2: Synthesize chapter abstracts (TIER_SPEED)
    Level 3: Audit argument continuity (TIER_LOGIC)
    """

    def __init__(self):
        """Initialize with the ModelRouter from LLM Gateway."""
        self.router = get_router()
        self._chunk_claims_cache: dict[str, list[ChunkClaim]] = {}

    def _level1_extract_claims(self, chunks: list[dict]) -> list[ChunkClaim]:
        """
        Level 1 (Map): Extract key claims from each chunk.

        Uses TIER_SPEED (Claude Haiku) for fast, cheap extraction.
        """
        logger.info(f"Level 1: Extracting claims from {len(chunks)} chunks")

        claims = []
        system_prompt = """You are a PhD thesis analyzer. Extract key claims and arguments.
Output JSON only with this structure:
{
    "key_claims": ["claim 1", "claim 2"],
    "methodology_notes": "any methodology mentioned or null",
    "evidence_cited": ["evidence 1", "evidence 2"]
}"""

        for chunk in chunks:
            prompt = f"""Extract the key academic claims from this thesis excerpt.
Focus on: main arguments, hypotheses, findings, and conclusions.

TEXT:
{chunk['text'][:6000]}"""  # Limit to avoid token issues

            try:
                response = self.router.generate(
                    prompt=prompt,
                    engine=EngineType.AIRLOCK,  # Routes to TIER_SPEED
                    system_prompt=system_prompt,
                )

                # Parse response
                content = response.content.strip()
                if content.startswith('```'):
                    content = re.sub(r'^```(?:json)?\n?', '', content)
                    content = re.sub(r'\n?```$', '', content)

                try:
                    data = json.loads(content)
                    claims.append(ChunkClaim(
                        chunk_index=chunk['index'],
                        key_claims=data.get('key_claims', []),
                        methodology_notes=data.get('methodology_notes'),
                        evidence_cited=data.get('evidence_cited', []),
                    ))
                except json.JSONDecodeError:
                    # Fallback: treat response as single claim
                    claims.append(ChunkClaim(
                        chunk_index=chunk['index'],
                        key_claims=[content[:500]],
                    ))

            except Exception as e:
                logger.warning(f"Level 1 extraction failed for chunk {chunk['index']}: {e}")
                claims.append(ChunkClaim(
                    chunk_index=chunk['index'],
                    key_claims=[f"[Extraction failed: {str(e)[:100]}]"],
                ))

        return claims

    def _level2_synthesize_chapters(
        self,
        chunks: list[dict],
        claims: list[ChunkClaim],
        chapter_boundaries: list[dict],
    ) -> list[ChapterAbstract]:
        """
        Level 2 (Reduce): Synthesize chapter abstracts from claims.

        Groups claims by chapter and creates consolidated abstracts.
        """
        logger.info(f"Level 2: Synthesizing {len(chapter_boundaries)} chapter abstracts")

        # If no chapter boundaries detected, create synthetic chapters
        if not chapter_boundaries:
            # Create 5 synthetic chapters based on position
            total_chunks = len(chunks)
            chunks_per_chapter = max(1, total_chunks // 5)
            chapter_boundaries = [
                {
                    'chapter_id': f'ch{i+1}',
                    'chapter_num': i + 1,
                    'title': ['Introduction', 'Literature Review', 'Methodology', 'Results', 'Conclusion'][i] if i < 5 else f'Section {i+1}',
                    'chunk_start': i * chunks_per_chapter,
                    'chunk_end': min((i + 1) * chunks_per_chapter, total_chunks),
                }
                for i in range(min(5, max(1, total_chunks // chunks_per_chapter + 1)))
            ]
        else:
            # Map chapters to chunk indices
            for ch in chapter_boundaries:
                ch['chunk_start'] = 0
                ch['chunk_end'] = len(chunks)
                for i, chunk in enumerate(chunks):
                    if chunk['start_char'] >= ch['start_char']:
                        ch['chunk_start'] = i
                        break
                for i, chunk in enumerate(chunks):
                    if chunk['end_char'] >= ch.get('end_char', float('inf')):
                        ch['chunk_end'] = i + 1
                        break

        abstracts = []
        system_prompt = """You are synthesizing chapter summaries for a PhD thesis.
Output JSON only:
{
    "core_argument": "The main thesis argument of this chapter",
    "key_claims": ["consolidated claim 1", "consolidated claim 2"],
    "research_question": "If chapter 1, state the research question, else null",
    "conclusion_answer": "If final chapter, state conclusion answer, else null"
}"""

        for ch in chapter_boundaries:
            # Gather claims for this chapter
            chapter_claims = [
                c for c in claims
                if ch.get('chunk_start', 0) <= c.chunk_index < ch.get('chunk_end', len(chunks))
            ]

            if not chapter_claims:
                continue

            # Flatten all claims for this chapter
            all_claims = []
            for cc in chapter_claims:
                all_claims.extend(cc.key_claims)

            claims_text = "\n".join(f"- {c}" for c in all_claims[:30])  # Limit

            prompt = f"""Synthesize these key claims into a chapter abstract.

CHAPTER: {ch['title']} (Chapter {ch.get('chapter_num', '?')})

KEY CLAIMS FROM THIS CHAPTER:
{claims_text}

Create a unified abstract capturing the core argument."""

            try:
                response = self.router.generate(
                    prompt=prompt,
                    engine=EngineType.AIRLOCK,  # TIER_SPEED
                    system_prompt=system_prompt,
                )

                content = response.content.strip()
                if content.startswith('```'):
                    content = re.sub(r'^```(?:json)?\n?', '', content)
                    content = re.sub(r'\n?```$', '', content)

                try:
                    data = json.loads(content)
                    abstracts.append(ChapterAbstract(
                        chapter_id=ch['chapter_id'],
                        chapter_title=ch['title'],
                        core_argument=data.get('core_argument', 'Not identified'),
                        key_claims=data.get('key_claims', all_claims[:5]),
                        research_question=data.get('research_question'),
                        conclusion_answer=data.get('conclusion_answer'),
                    ))
                except json.JSONDecodeError:
                    abstracts.append(ChapterAbstract(
                        chapter_id=ch['chapter_id'],
                        chapter_title=ch['title'],
                        core_argument=content[:500],
                        key_claims=all_claims[:5],
                    ))

            except Exception as e:
                logger.warning(f"Level 2 synthesis failed for {ch['chapter_id']}: {e}")
                abstracts.append(ChapterAbstract(
                    chapter_id=ch['chapter_id'],
                    chapter_title=ch['title'],
                    core_argument=f"[Synthesis failed: {str(e)[:100]}]",
                    key_claims=all_claims[:3],
                ))

        return abstracts

    def _level3_audit_continuity(
        self,
        chapter_abstracts: list[ChapterAbstract],
    ) -> dict:
        """
        Level 3 (Audit): Analyze argument continuity using TIER_LOGIC.

        Sends only chapter abstracts to Claude 3.5 Sonnet for deep analysis.
        """
        logger.info(f"Level 3: Auditing continuity across {len(chapter_abstracts)} chapters")

        if len(chapter_abstracts) < 2:
            return {
                'thread_status': 'insufficient_data',
                'continuity_score': 0,
                'missing_links': [],
                'analysis': 'Need at least 2 chapters to analyze continuity.',
                'visual_graph_nodes': [],
                'visual_graph_edges': [],
            }

        # Format chapter abstracts for the audit
        abstracts_text = ""
        for ab in chapter_abstracts:
            abstracts_text += f"""
### {ab.chapter_title} ({ab.chapter_id})
**Core Argument:** {ab.core_argument}
**Key Claims:** {', '.join(ab.key_claims[:5])}
"""
            if ab.research_question:
                abstracts_text += f"**Research Question:** {ab.research_question}\n"
            if ab.conclusion_answer:
                abstracts_text += f"**Conclusion Answer:** {ab.conclusion_answer}\n"

        system_prompt = """You are a PhD thesis examiner analyzing argument continuity.
Your task: Determine if the Conclusion adequately answers the Research Question.

You MUST respond with ONLY valid JSON in this exact structure:
{
    "thread_status": "solid" or "broken",
    "continuity_score": 0-100,
    "analysis": "2-3 paragraph detailed analysis",
    "missing_links": [
        {
            "from_chapter": "ch1",
            "to_chapter": "ch5",
            "description": "What logical connection is missing",
            "severity": "high/medium/low",
            "suggestion": "How to address this"
        }
    ],
    "visual_graph_nodes": [
        {
            "id": "node_1",
            "label": "Research Question",
            "type": "question",
            "chapter": "ch1"
        }
    ],
    "visual_graph_edges": [
        {
            "source": "node_1",
            "target": "node_2",
            "label": "leads to",
            "strength": 0.8
        }
    ]
}"""

        prompt = f"""Analyze the argument continuity of this PhD thesis.

CHAPTER ABSTRACTS:
{abstracts_text}

CRITICAL QUESTIONS:
1. Does the Conclusion (final chapter) answer the Research Question (Chapter 1)?
2. Is there a logical "red thread" connecting each chapter?
3. Are there missing logical links between chapters?
4. Create a visual graph showing the argument flow.

Respond with JSON only. No markdown formatting around the JSON."""

        try:
            response = self.router.generate(
                prompt=prompt,
                engine=EngineType.RED_THREAD,  # Routes to TIER_LOGIC (Claude 3.5 Sonnet)
                system_prompt=system_prompt,
            )

            content = response.content.strip()

            # Clean markdown wrappers
            if content.startswith('```'):
                content = re.sub(r'^```(?:json)?\n?', '', content)
                content = re.sub(r'\n?```$', '', content)

            result = json.loads(content)

            # Ensure required fields
            result.setdefault('thread_status', 'unknown')
            result.setdefault('continuity_score', 50)
            result.setdefault('missing_links', [])
            result.setdefault('analysis', 'Analysis completed.')
            result.setdefault('visual_graph_nodes', [])
            result.setdefault('visual_graph_edges', [])

            return result

        except json.JSONDecodeError as e:
            logger.error(f"Level 3 JSON parse error: {e}")
            return {
                'thread_status': 'error',
                'continuity_score': 0,
                'missing_links': [],
                'analysis': f'Failed to parse analysis: {str(e)}',
                'visual_graph_nodes': [],
                'visual_graph_edges': [],
                'raw_response': response.content if 'response' in dir() else None,
            }
        except Exception as e:
            logger.error(f"Level 3 audit error: {e}")
            return {
                'thread_status': 'error',
                'continuity_score': 0,
                'missing_links': [],
                'analysis': f'Audit failed: {str(e)}',
                'visual_graph_nodes': [],
                'visual_graph_edges': [],
            }

    def check_continuity(self, doc_text: str) -> RedThreadResult:
        """
        Main entry point: Run the full hierarchical logic check.

        Args:
            doc_text: Full thesis text to analyze

        Returns:
            RedThreadResult with continuity analysis
        """
        start_time = datetime.now()

        # Step 1: Chunk the text
        chunks = chunk_text_sliding_window(doc_text)
        logger.info(f"Created {len(chunks)} chunks from {len(doc_text)} chars")

        # Step 2: Identify chapter boundaries
        chapter_boundaries = identify_chapter_boundaries(doc_text)
        logger.info(f"Identified {len(chapter_boundaries)} chapter boundaries")

        # Step 3: Level 1 - Extract claims from chunks
        chunk_claims = self._level1_extract_claims(chunks)

        # Step 4: Level 2 - Synthesize chapter abstracts
        chapter_abstracts = self._level2_synthesize_chapters(
            chunks, chunk_claims, chapter_boundaries
        )

        # Step 5: Level 3 - Audit continuity
        audit_result = self._level3_audit_continuity(chapter_abstracts)

        # Build final result
        processing_time = (datetime.now() - start_time).total_seconds()

        # Determine human-readable status
        score = audit_result.get('continuity_score', 0)
        if score >= 80:
            status = "Strong Continuity"
        elif score >= 60:
            status = "Moderate Continuity"
        elif score >= 40:
            status = "Weak Continuity"
        else:
            status = "Broken Thread"

        # Parse missing links into Pydantic models
        missing_links = []
        for ml in audit_result.get('missing_links', []):
            try:
                missing_links.append(MissingLink(
                    from_chapter=ml.get('from_chapter', 'unknown'),
                    to_chapter=ml.get('to_chapter', 'unknown'),
                    description=ml.get('description', 'No description'),
                    severity=ml.get('severity', 'medium'),
                    suggestion=ml.get('suggestion', 'Review the connection'),
                ))
            except Exception:
                pass

        # Parse graph nodes
        graph_nodes = []
        for node in audit_result.get('visual_graph_nodes', []):
            try:
                graph_nodes.append(GraphNode(
                    id=node.get('id', f'node_{len(graph_nodes)}'),
                    label=node.get('label', 'Unknown'),
                    type=node.get('type', 'argument'),
                    chapter=node.get('chapter', 'unknown'),
                ))
            except Exception:
                pass

        # Parse graph edges
        graph_edges = []
        for edge in audit_result.get('visual_graph_edges', []):
            try:
                graph_edges.append(GraphEdge(
                    source=edge.get('source', ''),
                    target=edge.get('target', ''),
                    label=edge.get('label'),
                    strength=edge.get('strength', 1.0),
                ))
            except Exception:
                pass

        return RedThreadResult(
            continuity_score=score,
            thread_status=audit_result.get('thread_status', 'unknown'),
            status=status,
            analysis=audit_result.get('analysis', 'Analysis complete.'),
            missing_links=missing_links,
            chapter_abstracts=chapter_abstracts,
            visual_graph_nodes=graph_nodes,
            visual_graph_edges=graph_edges,
            processing_metadata={
                'total_chunks': len(chunks),
                'chapters_identified': len(chapter_boundaries),
                'claims_extracted': sum(len(c.key_claims) for c in chunk_claims),
                'processing_time_seconds': processing_time,
                'timestamp': datetime.now().isoformat(),
            },
        )


# =============================================================================
# Legacy RedThreadEngine (Backward Compatibility)
# =============================================================================

class RedThreadEngine:
    """
    Engine for maintaining logical continuity across thesis drafts.

    Now uses the HierarchicalLogicChecker for the check_continuity method
    while maintaining backward compatibility with vector-based methods.
    """

    def __init__(self, use_local: bool = False):
        """Initialize the Red Thread Engine."""
        # Initialize vector store for legacy operations
        if use_local:
            from core.vector_store import ChromaVectorStore
            self.vector_store = ChromaVectorStore(COLLECTION_NAME)
        else:
            self.vector_store = get_vector_store(COLLECTION_NAME)

        self.backend = self.vector_store.backend

        # Initialize the new hierarchical checker
        self._hierarchical_checker = HierarchicalLogicChecker()

        # Legacy Anthropic client for backward compatibility
        api_key = get_secret("ANTHROPIC_API_KEY")
        try:
            import anthropic
            self.claude_client = anthropic.Anthropic(api_key=api_key) if api_key else None
        except ImportError:
            self.claude_client = None

    def _extract_paragraphs(self, text: str, min_words: int = 20) -> list[str]:
        """Extract meaningful paragraphs from text."""
        paragraphs = re.split(r'\n\s*\n', text)
        cleaned = []
        for para in paragraphs:
            para = para.strip()
            if para and len(para.split()) >= min_words:
                cleaned.append(para)
        return cleaned

    def index_document(self, filepath: Path) -> int:
        """Index a single .docx document into the vector store."""
        try:
            doc = Document(filepath)
            full_text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            paragraphs = self._extract_paragraphs(full_text)

            if not paragraphs:
                return 0

            ids = [f"{filepath.stem}_{i}" for i in range(len(paragraphs))]
            metadatas = [
                {
                    "source_file": filepath.name,
                    "paragraph_index": i,
                    "word_count": len(p.split())
                }
                for i, p in enumerate(paragraphs)
            ]

            self.vector_store.upsert(
                ids=ids,
                documents=paragraphs,
                metadatas=metadatas
            )
            return len(paragraphs)
        except Exception as e:
            print(f"Error indexing {filepath.name}: {e}")
            return 0

    def index_drafts_folder(self, drafts_dir: Path = DRAFTS_DIR) -> dict:
        """Index all .docx files in the drafts folder."""
        stats = {
            "files_processed": 0,
            "paragraphs_indexed": 0,
            "files": [],
            "waiting_for_content": False
        }

        if not drafts_dir.exists():
            drafts_dir.mkdir(parents=True, exist_ok=True)
            stats["waiting_for_content"] = True
            return stats

        docx_files = list(drafts_dir.glob("*.docx"))

        if not docx_files:
            stats["waiting_for_content"] = True
            return stats

        for docx_file in docx_files:
            count = self.index_document(docx_file)
            stats["files_processed"] += 1
            stats["paragraphs_indexed"] += count
            stats["files"].append({"name": docx_file.name, "paragraphs": count})

        return stats

    def find_similar_passages(
        self,
        text: str,
        n_results: int = 5,
        threshold: float = 0.7
    ) -> list[dict]:
        """Find passages in the index similar to the given text."""
        if self.vector_store.count() == 0:
            return []

        results = self.vector_store.query(
            query_text=text,
            n_results=min(n_results, self.vector_store.count())
        )

        similar = []
        for i, (doc, meta, dist) in enumerate(zip(
            results["documents"][0],
            results["metadatas"][0],
            results["distances"][0]
        )):
            similarity = 1 / (1 + dist)
            if similarity >= threshold:
                similar.append({
                    "text": doc,
                    "source_file": meta["source_file"],
                    "paragraph_index": meta["paragraph_index"],
                    "similarity": round(similarity, 3),
                    "distance": round(dist, 3)
                })
        return similar

    def check_continuity(self, new_text: str) -> dict:
        """
        Check argument continuity using the Hierarchical Logic Check.

        This is the NEW implementation using Map-Reduce architecture.

        Args:
            new_text: The thesis text to analyze

        Returns:
            Dict compatible with frontend RedThreadModule.tsx
        """
        result = self._hierarchical_checker.check_continuity(new_text)

        # Return as dict for backward compatibility
        return result.model_dump()

    def get_stats(self) -> dict:
        """Get statistics about the current index."""
        stats = self.vector_store.get_stats()
        stats["total_paragraphs"] = self.vector_store.count()
        return stats

    def clear_index(self):
        """Clear all indexed content."""
        self.vector_store.delete_all()

    def index_existing_chapters(self, drafts_dir: Path = DRAFTS_DIR) -> dict:
        """Index all .docx files from /drafts folder into the vector store."""
        report = {
            "success": False,
            "timestamp": datetime.now().isoformat(),
            "total_files": 0,
            "total_paragraphs": 0,
            "total_words": 0,
            "chapters": [],
            "storage_path": str(CHROMA_DIR)
        }

        CHROMA_DIR.mkdir(parents=True, exist_ok=True)

        if not drafts_dir.exists():
            drafts_dir.mkdir(parents=True, exist_ok=True)
            report["error"] = f"Created empty drafts directory at {drafts_dir}"
            return report

        docx_files = list(drafts_dir.glob("*.docx"))

        if not docx_files:
            report["error"] = "No .docx files found in drafts folder"
            return report

        for docx_file in docx_files:
            try:
                doc = Document(docx_file)
                full_text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                paragraphs = self._extract_paragraphs(full_text, min_words=15)

                if not paragraphs:
                    continue

                word_count = sum(len(p.split()) for p in paragraphs)
                chapter_name = docx_file.stem
                ids = [f"{chapter_name}_para_{i}" for i in range(len(paragraphs))]

                metadatas = [
                    {
                        "source_file": docx_file.name,
                        "chapter": chapter_name,
                        "paragraph_index": i,
                        "word_count": len(p.split()),
                        "char_count": len(p),
                        "indexed_at": datetime.now().isoformat()
                    }
                    for i, p in enumerate(paragraphs)
                ]

                self.vector_store.upsert(ids=ids, documents=paragraphs, metadatas=metadatas)

                report["total_files"] += 1
                report["total_paragraphs"] += len(paragraphs)
                report["total_words"] += word_count
                report["chapters"].append({
                    "filename": docx_file.name,
                    "chapter": chapter_name,
                    "paragraphs": len(paragraphs),
                    "words": word_count
                })

            except Exception as e:
                report["chapters"].append({"filename": docx_file.name, "error": str(e)})

        report["success"] = report["total_paragraphs"] > 0
        return report

    def verify_consistency(self, new_draft_text: str) -> dict:
        """
        Verify consistency of new draft text.

        Now uses the hierarchical logic check internally.
        """
        return self.check_continuity(new_draft_text)

    def get_consistency_report_for_ui(self, new_draft_text: str) -> dict:
        """Get consistency report formatted for UI display."""
        full_report = self.check_continuity(new_draft_text)

        return {
            "status": full_report.get('status', 'unknown'),
            "score": full_report.get('continuity_score', 0),
            "score_label": self._score_to_label(full_report.get('continuity_score', 0)),
            "summary": full_report.get('analysis', ''),
            "thread_status": full_report.get('thread_status', 'unknown'),
            "missing_links": full_report.get('missing_links', []),
            "visual_graph": {
                "nodes": full_report.get('visual_graph_nodes', []),
                "edges": full_report.get('visual_graph_edges', []),
            },
            "full_report": full_report,
        }

    def _score_to_label(self, score: int) -> str:
        """Convert numeric score to human-readable label."""
        if score >= 95:
            return "Excellent"
        elif score >= 85:
            return "Good"
        elif score >= 70:
            return "Fair"
        elif score >= 50:
            return "Needs Review"
        else:
            return "Critical Issues"


# =============================================================================
# STANDALONE FUNCTIONS (for direct import)
# =============================================================================

def check_continuity(doc_text: str) -> dict:
    """
    Standalone function for hierarchical continuity check.

    Usage:
        from core.red_thread import check_continuity
        result = check_continuity(full_thesis_text)

    Returns:
        RedThreadResult as dict with thread_status, missing_links, visual_graph_nodes
    """
    checker = HierarchicalLogicChecker()
    result = checker.check_continuity(doc_text)
    return result.model_dump()


def index_existing_chapters(drafts_dir: Path = DRAFTS_DIR) -> dict:
    """Standalone function to index all chapters from /drafts folder."""
    engine = RedThreadEngine()
    return engine.index_existing_chapters(drafts_dir)


def verify_consistency(new_draft_text: str) -> dict:
    """Standalone function to verify consistency of new text."""
    engine = RedThreadEngine()
    return engine.verify_consistency(new_draft_text)


def main():
    """CLI for Red Thread Engine."""
    print("=" * 60)
    print("PHDx Red Thread Engine - Hierarchical Logic Check")
    print("=" * 60)

    engine = RedThreadEngine()

    # Index drafts
    print("\n[1] Indexing drafts folder...")
    stats = engine.index_drafts_folder()
    print(f"Indexed {stats['paragraphs_indexed']} paragraphs from {stats['files_processed']} files")

    # Show stats
    print("\n[2] Current index stats:")
    index_stats = engine.get_stats()
    print(f"  Total paragraphs: {index_stats['total_paragraphs']}")

    if index_stats['total_paragraphs'] > 0:
        print("\n[3] Ready for hierarchical continuity checks!")
        print("Use: result = check_continuity(thesis_text)")
    else:
        print("\n[3] No content indexed. Add .docx files to drafts/ folder.")


if __name__ == "__main__":
    main()
