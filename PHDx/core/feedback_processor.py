"""
Supervisor Feedback Processor - PHDx

Parses supervisor feedback documents (PDF/DOCX/TXT), categorizes feedback
using Claude AI with a Traffic Light System, and suggests revisions
that maintain the author's voice.

Traffic Light Categories:
🔴 Red: Critical structural/theoretical changes
🟡 Amber: Stylistic or citation-related corrections
🟢 Green: Positive reinforcement to be maintained
"""

import json
import os
import re
import hashlib
from datetime import datetime
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, asdict
from enum import Enum

import anthropic
from dotenv import load_dotenv

# Document parsing libraries
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Import ethics utilities and secrets
try:
    from core.ethics_utils import scrub_text, log_ai_usage
    from core.secrets_utils import get_secret
except ImportError:
    import sys
    sys.path.insert(0, str(Path(__file__).parent.parent))
    try:
        from core.ethics_utils import scrub_text, log_ai_usage
        from core.secrets_utils import get_secret
    except ImportError:
        def scrub_text(text):
            return {"scrubbed_text": text, "total_redactions": 0}
        def log_ai_usage(*args, **kwargs):
            pass
        def get_secret(key, default=None):
            return os.getenv(key, default)

load_dotenv()

# Paths
ROOT_DIR = Path(__file__).parent.parent
FEEDBACK_DIR = ROOT_DIR / "feedback"
DRAFTS_DIR = ROOT_DIR / "drafts"
DATA_DIR = ROOT_DIR / "data"
FEEDBACK_CACHE = DATA_DIR / "feedback_analysis.json"
DNA_PATH = DATA_DIR / "author_dna.json"


class TrafficLight(Enum):
    """Traffic Light categorization for feedback severity."""
    RED = "red"      # 🔴 Critical structural/theoretical changes
    AMBER = "amber"  # 🟡 Stylistic or citation-related corrections
    GREEN = "green"  # 🟢 Positive reinforcement to be maintained


class FeedbackCategory(Enum):
    """Detailed categories for supervisor feedback."""
    MAJOR_STRUCTURAL = "major_structural"
    THEORETICAL = "theoretical"
    MINOR_STYLISTIC = "minor_stylistic"
    CITATIONS_NEEDED = "citations_needed"
    POSITIVE = "positive"
    GENERAL = "general"


@dataclass
class FeedbackItem:
    """A single piece of categorized feedback with Traffic Light status."""
    id: str
    text: str
    category: str
    traffic_light: str  # red, amber, green
    priority: str  # high, medium, low
    chapter: str
    section: str
    target_paragraph: str  # text snippet to highlight
    action_required: str
    resolved: bool = False
    source_file: str = ""
    created_at: str = ""
    suggested_revision: str = ""  # AI-suggested revision maintaining author voice

    def to_dict(self) -> dict:
        return asdict(self)

    @classmethod
    def from_dict(cls, data: dict) -> "FeedbackItem":
        # Handle backwards compatibility for items without traffic_light
        if "traffic_light" not in data:
            # Map old categories to traffic lights
            cat = data.get("category", "general")
            if cat in ["major_structural", "theoretical"]:
                data["traffic_light"] = "red"
            elif cat in ["minor_stylistic", "citations_needed"]:
                data["traffic_light"] = "amber"
            elif cat == "positive":
                data["traffic_light"] = "green"
            else:
                data["traffic_light"] = "amber"
        if "suggested_revision" not in data:
            data["suggested_revision"] = ""
        return cls(**data)

    @property
    def traffic_light_emoji(self) -> str:
        """Get the emoji for this item's traffic light status."""
        return {"red": "🔴", "amber": "🟡", "green": "🟢"}.get(self.traffic_light, "⚪")


def load_author_dna() -> Optional[dict]:
    """Load the author's DNA profile for voice matching."""
    if DNA_PATH.exists():
        try:
            with open(DNA_PATH, 'r', encoding='utf-8') as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            return None
    return None


def suggest_revision(
    feedback_snippet: str,
    original_text: str,
    author_dna: Optional[dict] = None,
    claude_client: Optional[anthropic.Anthropic] = None
) -> str:
    """
    Generate a revision suggestion that addresses supervisor feedback
    while maintaining the author's unique writing voice.

    Args:
        feedback_snippet: The supervisor's feedback/critique
        original_text: The original text that needs revision
        author_dna: The author's linguistic fingerprint (from author_dna.json)
        claude_client: Anthropic client instance

    Returns:
        Suggested revision text in the author's voice
    """
    if not claude_client:
        api_key = get_secret("ANTHROPIC_API_KEY")
        if not api_key:
            return "[Revision unavailable - ANTHROPIC_API_KEY not configured]"
        claude_client = anthropic.Anthropic(api_key=api_key)

    if not author_dna:
        author_dna = load_author_dna()

    # Build DNA context for voice matching
    dna_context = ""
    if author_dna:
        sentence = author_dna.get("sentence_complexity", {})
        hedging = author_dna.get("hedging_analysis", {})
        transitions = author_dna.get("transition_vocabulary", {})

        dna_context = f"""
AUTHOR'S WRITING DNA (maintain this voice):
- Average sentence length: {sentence.get('average_length', 20)} words
- Hedging density: {hedging.get('hedging_density_per_1000_words', 5)} per 1000 words
- Top hedging phrases: {', '.join(list(hedging.get('phrases_found', {}).keys())[:5])}
- Preferred transitions: {', '.join(transitions.get('preferred_categories', ['contrast', 'addition'])[:3])}
- Sentence style: Mix of {sentence.get('length_distribution', {})}
"""

    prompt = f"""You are a PhD thesis revision assistant. Your task is to rewrite a section of text to address supervisor feedback while PERFECTLY maintaining the author's unique writing voice.

{dna_context}

SUPERVISOR'S FEEDBACK:
{feedback_snippet}

ORIGINAL TEXT TO REVISE:
{original_text}

INSTRUCTIONS:
1. Address the supervisor's critique directly
2. Maintain the author's characteristic sentence length and complexity
3. Use the author's preferred hedging language (e.g., {', '.join(list(author_dna.get('hedging_analysis', {}).get('phrases_found', {'may': 1, 'could': 1}).keys())[:3]) if author_dna else 'may, could, arguably'})
4. Employ the author's transition vocabulary style
5. Keep the academic tone consistent with the original
6. Do NOT add commentary - output ONLY the revised text

Write the revised passage now:"""

    # Log AI usage
    log_ai_usage(
        action_type="suggest_revision",
        data_source="feedback_processor",
        prompt=f"Revision for: {feedback_snippet[:100]}...",
        was_scrubbed=False,
        redactions_count=0
    )

    try:
        response = claude_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2048,
            messages=[{"role": "user", "content": prompt}]
        )
        return response.content[0].text.strip()
    except Exception as e:
        return f"[Revision error: {e}]"


class FeedbackProcessor:
    """
    Processes supervisor feedback documents and categorizes comments.

    Features:
    - Parses PDF, DOCX, and TXT files
    - Uses Claude to categorize feedback
    - Maps feedback to thesis sections
    - Tracks resolution status
    """

    def __init__(self):
        """Initialize the feedback processor."""
        self.feedback_items: list[FeedbackItem] = []
        self.processed_files: dict[str, str] = {}  # filename -> hash

        # Initialize Claude client
        api_key = get_secret("ANTHROPIC_API_KEY")
        self.claude = anthropic.Anthropic(api_key=api_key) if api_key else None

        # Load existing feedback
        self._load_cache()

    def _load_cache(self):
        """Load cached feedback analysis."""
        if FEEDBACK_CACHE.exists():
            try:
                with open(FEEDBACK_CACHE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.feedback_items = [
                        FeedbackItem.from_dict(item)
                        for item in data.get("items", [])
                    ]
                    self.processed_files = data.get("processed_files", {})
            except (json.JSONDecodeError, IOError):
                self.feedback_items = []

    def _save_cache(self):
        """Save feedback analysis to cache."""
        DATA_DIR.mkdir(parents=True, exist_ok=True)
        with open(FEEDBACK_CACHE, 'w', encoding='utf-8') as f:
            json.dump({
                "items": [item.to_dict() for item in self.feedback_items],
                "processed_files": self.processed_files,
                "last_updated": datetime.now().isoformat()
            }, f, indent=2)

    def _get_file_hash(self, filepath: Path) -> str:
        """Get hash of file for change detection."""
        with open(filepath, 'rb') as f:
            return hashlib.md5(f.read(), usedforsecurity=False).hexdigest()

    def _parse_txt(self, filepath: Path) -> str:
        """Parse a text file."""
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _parse_docx(self, filepath: Path) -> str:
        """Parse a DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc = DocxDocument(filepath)
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract comments if present
        # Note: python-docx doesn't easily extract comments,
        # so we focus on the main text

        return "\n\n".join(paragraphs)

    def _parse_pdf(self, filepath: Path) -> str:
        """Parse a PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PyMuPDF not installed. Run: pip install pymupdf")

        doc = fitz.open(filepath)
        text_parts = []

        for page in doc:
            text_parts.append(page.get_text())

        doc.close()
        return "\n\n".join(text_parts)

    def parse_document(self, filepath: Path) -> str:
        """
        Parse a document based on its extension.

        Supports: .txt, .md, .docx, .pdf
        """
        suffix = filepath.suffix.lower()

        if suffix in ['.txt', '.md']:
            return self._parse_txt(filepath)
        elif suffix == '.docx':
            return self._parse_docx(filepath)
        elif suffix == '.pdf':
            return self._parse_pdf(filepath)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _get_draft_context(self) -> str:
        """Get summary of draft structure for context."""
        context_parts = []

        if not DRAFTS_DIR.exists():
            return "No drafts available."

        for docx_file in DRAFTS_DIR.glob("*.docx"):
            try:
                if DOCX_AVAILABLE:
                    doc = DocxDocument(docx_file)
                    # Get first few paragraphs as context
                    preview = []
                    for para in doc.paragraphs[:10]:
                        if para.text.strip():
                            preview.append(para.text.strip()[:200])

                    context_parts.append(f"**{docx_file.name}**:\n" + "\n".join(preview[:5]))
            except Exception:
                continue

        return "\n\n---\n\n".join(context_parts) if context_parts else "No drafts available."

    def categorize_feedback(self, feedback_text: str, source_file: str = "") -> list[FeedbackItem]:
        """
        Use Claude to categorize feedback into structured items with Traffic Light system.

        Traffic Light System:
        🔴 Red: Critical structural/theoretical changes
        🟡 Amber: Stylistic or citation-related corrections
        🟢 Green: Positive reinforcement to be maintained

        Args:
            feedback_text: Raw text from feedback document
            source_file: Name of the source file

        Returns:
            List of categorized FeedbackItem objects with traffic light status
        """
        if not self.claude:
            # Return uncategorized if no Claude available
            return [FeedbackItem(
                id=hashlib.md5(feedback_text[:100].encode(), usedforsecurity=False).hexdigest()[:8],
                text=feedback_text,
                category=FeedbackCategory.GENERAL.value,
                traffic_light="amber",
                priority="medium",
                chapter="Unknown",
                section="Unknown",
                target_paragraph="",
                action_required="Review feedback",
                source_file=source_file,
                created_at=datetime.now().isoformat()
            )]

        # Get draft context for mapping
        draft_context = self._get_draft_context()

        # Scrub feedback before sending to AI
        scrub_result = scrub_text(feedback_text)
        scrubbed_feedback = scrub_result["scrubbed_text"]

        prompt = f"""You are a PhD thesis feedback analyzer. Parse the supervisor feedback and categorize each distinct comment using a TRAFFIC LIGHT system.

FEEDBACK DOCUMENT:
{scrubbed_feedback[:8000]}

DRAFT CONTEXT (for mapping feedback to sections):
{draft_context[:3000]}

For each distinct piece of feedback, create a JSON object with:

- "text": The feedback comment (verbatim or paraphrased)

- "traffic_light": CRITICAL - Use this Traffic Light system:
  * "red": 🔴 Critical structural/theoretical changes - fundamental issues with argument, methodology, or theory
  * "amber": 🟡 Stylistic or citation-related corrections - grammar, clarity, formatting, missing references
  * "green": 🟢 Positive reinforcement - things done well that should be maintained

- "category": One of:
  * "major_structural": Chapter organization, argument flow, missing sections
  * "theoretical": Theoretical framework issues, conceptual problems
  * "minor_stylistic": Grammar, clarity, word choice, formatting
  * "citations_needed": Missing references, citation format issues
  * "positive": Praise, things to maintain
  * "general": Other feedback

- "priority": "high", "medium", or "low" based on impact
- "chapter": Which chapter this applies to (e.g., "Chapter 1", "Literature Review")
- "section": Specific section if identifiable
- "target_paragraph": A short text snippet (10-20 words) from the draft for highlighting
- "action_required": Brief action statement (e.g., "Restructure argument", "Add citation")

IMPORTANT: Include positive feedback (green) as well as critiques!

Return a JSON array. Be thorough - extract ALL feedback points.

Example:
[
  {{
    "text": "The theoretical framework lacks engagement with postcolonial perspectives",
    "traffic_light": "red",
    "category": "theoretical",
    "priority": "high",
    "chapter": "Chapter 2",
    "section": "2.3 Theoretical Framework",
    "target_paragraph": "This thesis adopts a critical realist ontology",
    "action_required": "Integrate postcolonial theoretical perspectives"
  }},
  {{
    "text": "Excellent use of hedging language throughout",
    "traffic_light": "green",
    "category": "positive",
    "priority": "low",
    "chapter": "General",
    "section": "",
    "target_paragraph": "",
    "action_required": "Maintain current hedging style"
  }}
]

Return ONLY valid JSON array, no markdown."""

        # Log AI usage
        log_ai_usage(
            action_type="categorize_feedback",
            data_source="feedback_document",
            prompt=f"Categorizing feedback from: {source_file}",
            was_scrubbed=scrub_result["total_redactions"] > 0,
            redactions_count=scrub_result["total_redactions"]
        )

        try:
            response = self.claude.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=4096,
                messages=[{"role": "user", "content": prompt}]
            )

            response_text = response.content[0].text.strip()

            # Clean markdown if present
            if response_text.startswith("```"):
                response_text = re.sub(r'^```(?:json)?\n?', '', response_text)
                response_text = re.sub(r'\n?```$', '', response_text)

            items_data = json.loads(response_text)

            feedback_items = []
            for i, item in enumerate(items_data):
                # Map category to traffic light if not provided
                traffic_light = item.get("traffic_light", "amber")
                if traffic_light not in ["red", "amber", "green"]:
                    cat = item.get("category", "general")
                    if cat in ["major_structural", "theoretical"]:
                        traffic_light = "red"
                    elif cat == "positive":
                        traffic_light = "green"
                    else:
                        traffic_light = "amber"

                feedback_items.append(FeedbackItem(
                    id=f"{hashlib.md5(source_file.encode(), usedforsecurity=False).hexdigest()[:4]}_{i:03d}",
                    text=item.get("text", ""),
                    category=item.get("category", "general"),
                    traffic_light=traffic_light,
                    priority=item.get("priority", "medium"),
                    chapter=item.get("chapter", "Unknown"),
                    section=item.get("section", ""),
                    target_paragraph=item.get("target_paragraph", ""),
                    action_required=item.get("action_required", ""),
                    source_file=source_file,
                    created_at=datetime.now().isoformat()
                ))

            return feedback_items

        except Exception as e:
            print(f"Error categorizing feedback: {e}")
            return [FeedbackItem(
                id=hashlib.md5(feedback_text[:100].encode(), usedforsecurity=False).hexdigest()[:8],
                text=feedback_text[:500],
                category=FeedbackCategory.GENERAL.value,
                traffic_light="amber",
                priority="medium",
                chapter="Unknown",
                section="Unknown",
                target_paragraph="",
                action_required="Review feedback manually",
                source_file=source_file,
                created_at=datetime.now().isoformat()
            )]

    def process_feedback_folder(self, force_reprocess: bool = False) -> dict:
        """
        Process all feedback documents in the /feedback folder.

        Args:
            force_reprocess: If True, reprocess all files even if cached

        Returns:
            Summary of processing results
        """
        if not FEEDBACK_DIR.exists():
            FEEDBACK_DIR.mkdir(parents=True)
            return {"status": "created", "message": "Feedback folder created", "items": 0}

        results = {
            "files_processed": 0,
            "files_skipped": 0,
            "new_items": 0,
            "errors": []
        }

        supported_extensions = {'.txt', '.md', '.docx', '.pdf'}

        for filepath in FEEDBACK_DIR.iterdir():
            if filepath.suffix.lower() not in supported_extensions:
                continue

            if filepath.name.startswith('.'):
                continue

            try:
                # Check if file has changed
                file_hash = self._get_file_hash(filepath)

                if not force_reprocess and filepath.name in self.processed_files:
                    if self.processed_files[filepath.name] == file_hash:
                        results["files_skipped"] += 1
                        continue

                # Parse and categorize
                feedback_text = self.parse_document(filepath)

                if not feedback_text.strip():
                    results["errors"].append(f"{filepath.name}: Empty document")
                    continue

                # Remove old items from this file
                self.feedback_items = [
                    item for item in self.feedback_items
                    if item.source_file != filepath.name
                ]

                # Categorize new feedback
                new_items = self.categorize_feedback(feedback_text, filepath.name)
                self.feedback_items.extend(new_items)

                # Update processed files
                self.processed_files[filepath.name] = file_hash

                results["files_processed"] += 1
                results["new_items"] += len(new_items)

            except Exception as e:
                results["errors"].append(f"{filepath.name}: {str(e)}")

        # Save to cache
        self._save_cache()

        return results

    def get_feedback_by_category(self) -> dict[str, list[FeedbackItem]]:
        """Get feedback items grouped by category."""
        grouped = {
            "major_structural": [],
            "theoretical": [],
            "minor_stylistic": [],
            "citations_needed": [],
            "positive": [],
            "general": []
        }

        for item in self.feedback_items:
            category = item.category
            if category in grouped:
                grouped[category].append(item)
            else:
                grouped["general"].append(item)

        return grouped

    def get_feedback_by_traffic_light(self) -> dict[str, list[FeedbackItem]]:
        """Get feedback items grouped by Traffic Light status."""
        grouped = {
            "red": [],    # 🔴 Critical
            "amber": [],  # 🟡 Corrections
            "green": []   # 🟢 Positive
        }

        for item in self.feedback_items:
            light = item.traffic_light
            if light in grouped:
                grouped[light].append(item)
            else:
                grouped["amber"].append(item)

        # Sort each group by priority
        priority_order = {"high": 0, "medium": 1, "low": 2}
        for light in grouped:
            grouped[light].sort(key=lambda x: (x.resolved, priority_order.get(x.priority, 1)))

        return grouped

    def get_traffic_light_counts(self) -> dict:
        """Get count of items by traffic light status."""
        counts = {"red": 0, "amber": 0, "green": 0, "total": 0}
        unresolved = {"red": 0, "amber": 0, "green": 0, "total": 0}

        for item in self.feedback_items:
            light = item.traffic_light if item.traffic_light in counts else "amber"
            counts[light] += 1
            counts["total"] += 1

            if not item.resolved:
                unresolved[light] += 1
                unresolved["total"] += 1

        return {"counts": counts, "unresolved": unresolved}

    def get_feedback_by_chapter(self) -> dict[str, list[FeedbackItem]]:
        """Get feedback items grouped by chapter."""
        grouped = {}

        for item in self.feedback_items:
            chapter = item.chapter or "Unknown"
            if chapter not in grouped:
                grouped[chapter] = []
            grouped[chapter].append(item)

        return grouped

    def get_unresolved_count(self) -> dict:
        """Get count of unresolved feedback by category."""
        counts = {
            "major_structural": 0,
            "minor_stylistic": 0,
            "citations_needed": 0,
            "general": 0,
            "total": 0
        }

        for item in self.feedback_items:
            if not item.resolved:
                category = item.category if item.category in counts else "general"
                counts[category] += 1
                counts["total"] += 1

        return counts

    def mark_resolved(self, item_id: str, resolved: bool = True) -> bool:
        """Mark a feedback item as resolved/unresolved."""
        for item in self.feedback_items:
            if item.id == item_id:
                item.resolved = resolved
                self._save_cache()
                return True
        return False

    def get_item_by_id(self, item_id: str) -> Optional[FeedbackItem]:
        """Get a specific feedback item by ID."""
        for item in self.feedback_items:
            if item.id == item_id:
                return item
        return None

    def get_stats(self) -> dict:
        """Get processor statistics."""
        by_category = self.get_feedback_by_category()
        unresolved = self.get_unresolved_count()

        return {
            "total_items": len(self.feedback_items),
            "files_processed": len(self.processed_files),
            "by_category": {k: len(v) for k, v in by_category.items()},
            "unresolved": unresolved,
            "resolved": len(self.feedback_items) - unresolved["total"]
        }


# =============================================================================
# STREAMLIT UI COMPONENTS
# =============================================================================

def render_feedback_tab(processor: FeedbackProcessor):
    """
    Render the Supervisor Feedback tab with Traffic Light Correction Checklist.

    Features:
    - Traffic Light categorization (🔴 Red, 🟡 Amber, 🟢 Green)
    - Interactive correction checklist
    - Suggest Revision button using author DNA
    - Click to highlight functionality
    """
    import streamlit as st

    st.markdown("### 🚦 Supervisor Feedback - Correction Checklist")
    st.markdown("*Traffic Light System: 🔴 Critical | 🟡 Corrections | 🟢 Positive*")

    # Process feedback button row
    col1, col2, col3 = st.columns([1, 1, 2])

    with col1:
        if st.button("📥 Process Feedback", type="primary", use_container_width=True):
            with st.spinner("Processing feedback documents..."):
                results = processor.process_feedback_folder()

                if results.get("errors"):
                    for err in results["errors"]:
                        st.error(err)

                st.success(f"Processed {results['files_processed']} files, {results['new_items']} items found")
                st.rerun()

    with col2:
        if st.button("🔄 Refresh", use_container_width=True):
            results = processor.process_feedback_folder(force_reprocess=True)
            st.rerun()

    with col3:
        stats = processor.get_stats()
        total = stats["total_items"]
        resolved = stats["resolved"]

        if total > 0:
            progress = resolved / total
            st.progress(progress, text=f"Progress: {resolved}/{total} resolved ({progress*100:.0f}%)")
        else:
            st.info("Add .txt or .docx files to /feedback folder")

    # Traffic Light Summary
    if total > 0:
        tl_stats = processor.get_traffic_light_counts()
        col1, col2, col3, col4 = st.columns(4)

        with col1:
            red_count = tl_stats["unresolved"]["red"]
            st.markdown(f"""
            <div style="text-align: center; padding: 0.5rem; background: rgba(244, 67, 54, 0.15);
                        border: 1px solid rgba(244, 67, 54, 0.4); border-radius: 8px;">
                <div style="font-size: 1.5rem;">🔴</div>
                <div style="font-size: 1.2rem; font-weight: bold; color: #f44336;">{red_count}</div>
                <div style="font-size: 0.75rem; color: rgba(224, 224, 224, 0.7);">Critical</div>
            </div>
            """, unsafe_allow_html=True)

        with col2:
            amber_count = tl_stats["unresolved"]["amber"]
            st.markdown(f"""
            <div style="text-align: center; padding: 0.5rem; background: rgba(255, 193, 7, 0.15);
                        border: 1px solid rgba(255, 193, 7, 0.4); border-radius: 8px;">
                <div style="font-size: 1.5rem;">🟡</div>
                <div style="font-size: 1.2rem; font-weight: bold; color: #ffc107;">{amber_count}</div>
                <div style="font-size: 0.75rem; color: rgba(224, 224, 224, 0.7);">Corrections</div>
            </div>
            """, unsafe_allow_html=True)

        with col3:
            green_count = tl_stats["counts"]["green"]
            st.markdown(f"""
            <div style="text-align: center; padding: 0.5rem; background: rgba(76, 175, 80, 0.15);
                        border: 1px solid rgba(76, 175, 80, 0.4); border-radius: 8px;">
                <div style="font-size: 1.5rem;">🟢</div>
                <div style="font-size: 1.2rem; font-weight: bold; color: #4caf50;">{green_count}</div>
                <div style="font-size: 0.75rem; color: rgba(224, 224, 224, 0.7);">Positive</div>
            </div>
            """, unsafe_allow_html=True)

        with col4:
            st.markdown(f"""
            <div style="text-align: center; padding: 0.5rem; background: rgba(0, 113, 206, 0.15);
                        border: 1px solid rgba(0, 113, 206, 0.4); border-radius: 8px;">
                <div style="font-size: 1.5rem;">📊</div>
                <div style="font-size: 1.2rem; font-weight: bold; color: #0071ce;">{resolved}/{total}</div>
                <div style="font-size: 0.75rem; color: rgba(224, 224, 224, 0.7);">Resolved</div>
            </div>
            """, unsafe_allow_html=True)

    # Initialize session state
    if "highlight_text" not in st.session_state:
        st.session_state.highlight_text = None
    if "highlight_chapter" not in st.session_state:
        st.session_state.highlight_chapter = None
    if "show_revision" not in st.session_state:
        st.session_state.show_revision = None

    # Show highlighted text notification
    if st.session_state.highlight_text:
        st.markdown(f"""
        <div style="background: rgba(0, 113, 206, 0.15); border: 1px solid rgba(0, 113, 206, 0.4);
                    border-radius: 8px; padding: 0.75rem; margin: 1rem 0;">
            <strong>🎯 Highlighting in Drafting Pane:</strong><br>
            <em>"{st.session_state.highlight_text[:100]}..."</em>
        </div>
        """, unsafe_allow_html=True)

        if st.button("Clear Highlight", key="clear_hl_feedback"):
            st.session_state.highlight_text = None
            st.session_state.highlight_chapter = None
            st.rerun()

    st.markdown("---")

    # Traffic Light tabs
    feedback_by_tl = processor.get_feedback_by_traffic_light()

    tl_tabs = st.tabs([
        f"🔴 Critical ({len(feedback_by_tl['red'])})",
        f"🟡 Corrections ({len(feedback_by_tl['amber'])})",
        f"🟢 Positive ({len(feedback_by_tl['green'])})"
    ])

    traffic_light_keys = ["red", "amber", "green"]
    traffic_light_styles = {
        "red": {"bg": "rgba(244, 67, 54, 0.1)", "border": "#f44336", "icon": "🔴"},
        "amber": {"bg": "rgba(255, 193, 7, 0.1)", "border": "#ffc107", "icon": "🟡"},
        "green": {"bg": "rgba(76, 175, 80, 0.1)", "border": "#4caf50", "icon": "🟢"}
    }
    priority_colors = {"high": "#f44336", "medium": "#ffc107", "low": "#4caf50"}

    for tab, tl_key in zip(tl_tabs, traffic_light_keys):
        with tab:
            items = feedback_by_tl[tl_key]
            style = traffic_light_styles[tl_key]

            if not items:
                if tl_key == "green":
                    st.info("No positive feedback recorded yet.")
                else:
                    st.success(f"No {tl_key} items remaining! 🎉")
                continue

            for item in items:
                checkbox_key = f"tl_{item.id}"
                priority_color = priority_colors.get(item.priority, "#ffc107")

                # Card container
                resolved_style = "opacity: 0.5;" if item.resolved else ""
                st.markdown(f"""
                <div style="background: {style['bg']}; border-left: 4px solid {style['border']};
                            border-radius: 8px; padding: 0.75rem; margin-bottom: 0.75rem; {resolved_style}">
                """, unsafe_allow_html=True)

                col1, col2, col3 = st.columns([0.5, 4, 1.5])

                with col1:
                    resolved = st.checkbox(
                        "",
                        value=item.resolved,
                        key=checkbox_key,
                        label_visibility="collapsed"
                    )
                    if resolved != item.resolved:
                        processor.mark_resolved(item.id, resolved)
                        st.rerun()

                with col2:
                    text_style = "text-decoration: line-through;" if item.resolved else ""

                    st.markdown(f"""
                    <div style="{text_style}">
                        <span style="color: {priority_color}; font-weight: bold; font-size: 0.8rem;">
                            [{item.priority.upper()}]
                        </span>
                        <strong>{item.chapter}</strong>
                        {f'- {item.section}' if item.section else ''}
                        <br>
                        <span style="color: rgba(224, 224, 224, 0.9);">
                            {item.text[:200]}{'...' if len(item.text) > 200 else ''}
                        </span>
                        <br>
                        <span style="color: {style['border']}; font-size: 0.85rem;">
                            ➜ {item.action_required}
                        </span>
                    </div>
                    """, unsafe_allow_html=True)

                with col3:
                    # Action buttons
                    if not item.resolved and tl_key != "green":
                        if item.target_paragraph:
                            if st.button("🎯", key=f"hl_{item.id}", help="Highlight in draft"):
                                st.session_state.highlight_text = item.target_paragraph
                                st.session_state.highlight_chapter = item.chapter
                                st.rerun()

                        if st.button("✨", key=f"rev_{item.id}", help="Suggest revision"):
                            st.session_state.show_revision = item.id

                st.markdown("</div>", unsafe_allow_html=True)

                # Show revision suggestion if requested
                if st.session_state.show_revision == item.id:
                    with st.expander("✨ Suggested Revision", expanded=True):
                        st.markdown("**Original feedback:**")
                        st.info(item.text)

                        if item.target_paragraph:
                            st.markdown("**Target text:**")
                            st.code(item.target_paragraph, language=None)

                            if st.button("Generate Revision", key=f"gen_{item.id}", type="primary"):
                                with st.spinner("Generating revision in your voice..."):
                                    revision = suggest_revision(
                                        item.text,
                                        item.target_paragraph,
                                        load_author_dna()
                                    )
                                    st.markdown("**Suggested revision:**")
                                    st.success(revision)

                        if st.button("Close", key=f"close_{item.id}"):
                            st.session_state.show_revision = None
                            st.rerun()

    # Summary by chapter
    st.markdown("---")
    st.markdown("#### 📚 Summary by Chapter")

    by_chapter = processor.get_feedback_by_chapter()

    if by_chapter:
        for chapter, items in sorted(by_chapter.items()):
            unresolved_count = sum(1 for i in items if not i.resolved)
            red_count = sum(1 for i in items if i.traffic_light == "red" and not i.resolved)
            amber_count = sum(1 for i in items if i.traffic_light == "amber" and not i.resolved)
            green_count = sum(1 for i in items if i.traffic_light == "green")

            if unresolved_count > 0:
                st.markdown(
                    f"**{chapter}**: "
                    f"{'🔴' * red_count}{'🟡' * amber_count}{'🟢' * green_count} "
                    f"({unresolved_count} unresolved)"
                )
            else:
                st.markdown(f"**{chapter}**: ✅ All resolved")
    else:
        st.info("Process feedback to see chapter summary.")


def get_highlight_text() -> Optional[str]:
    """
    Get the current text to highlight in the drafting pane.

    Call this from the drafting tab to check if text should be highlighted.
    """
    import streamlit as st
    return st.session_state.get("highlight_text")


def get_highlight_chapter() -> Optional[str]:
    """Get the chapter associated with the current highlight."""
    import streamlit as st
    return st.session_state.get("highlight_chapter")


# =============================================================================
# CLI
# =============================================================================

def main():
    """CLI for testing feedback processor."""
    print("=" * 60)
    print("PHDx Supervisor Feedback Processor")
    print("=" * 60)

    processor = FeedbackProcessor()

    print(f"\nFeedback directory: {FEEDBACK_DIR}")
    print(f"PDF support: {'Yes' if PDF_AVAILABLE else 'No (pip install pymupdf)'}")
    print(f"DOCX support: {'Yes' if DOCX_AVAILABLE else 'No (pip install python-docx)'}")
    print(f"Claude API: {'Configured' if processor.claude else 'Not configured'}")

    # Process feedback folder
    print("\nProcessing feedback folder...")
    results = processor.process_feedback_folder()

    print("\nResults:")
    print(f"  Files processed: {results['files_processed']}")
    print(f"  Files skipped: {results['files_skipped']}")
    print(f"  New items: {results['new_items']}")

    if results['errors']:
        print("\n  Errors:")
        for err in results['errors']:
            print(f"    - {err}")

    # Show stats
    stats = processor.get_stats()
    print("\nFeedback Statistics:")
    print(f"  Total items: {stats['total_items']}")
    print(f"  Resolved: {stats['resolved']}")
    print(f"  Unresolved: {stats['unresolved']['total']}")
    print("\n  By category:")
    for cat, count in stats['by_category'].items():
        print(f"    - {cat}: {count}")

    # Show sample items
    if processor.feedback_items:
        print("\nSample feedback items:")
        for item in processor.feedback_items[:3]:
            print(f"\n  [{item.priority.upper()}] {item.category}")
            print(f"  Chapter: {item.chapter}")
            print(f"  Text: {item.text[:80]}...")
            print(f"  Action: {item.action_required}")


if __name__ == "__main__":
    main()
