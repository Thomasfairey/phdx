#!/usr/bin/env python3
"""
Synthetic Research Environment Generator for PHDx Testing

Creates realistic PhD thesis drafts, mock data, and supervisor feedback
for comprehensive system testing.
"""

import csv
import random
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
ROOT_DIR = Path(__file__).parent
DRAFTS_DIR = ROOT_DIR / "drafts"
DATA_DIR = ROOT_DIR / "data"
FEEDBACK_DIR = ROOT_DIR / "feedback"

# Ensure directories exist
DRAFTS_DIR.mkdir(parents=True, exist_ok=True)
DATA_DIR.mkdir(parents=True, exist_ok=True)
FEEDBACK_DIR.mkdir(parents=True, exist_ok=True)


def create_synthetic_intro():
    """
    Generate synthetic_intro.docx: 1,000 words on 'The Digital Panopticon in Urban Governance'
    with formal, hedging academic tone.
    """
    doc = Document()

    # Title
    title = doc.add_heading("Chapter 1: Introduction", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("The Digital Panopticon in Urban Governance", level=2)

    content = """
The proliferation of digital surveillance technologies within contemporary urban environments arguably represents one of the most significant transformations in the relationship between state apparatus and citizen subjectivity. This thesis contends that the emergence of what may be termed the 'digital panopticon' fundamentally reconfigures the modalities through which governmental power operates, particularly within the context of smart city initiatives. While scholars have extensively examined the implications of surveillance capitalism in commercial contexts, there appears to be a relative paucity of critical engagement with the ways in which municipal governance has appropriated and adapted these technological affordances for ostensibly public purposes.

The conceptual framework underpinning this investigation draws substantially upon Foucauldian analytics of power, particularly the notion of disciplinary mechanisms that operate through visibility and normalisation. However, it would seem reductive to simply transpose Bentham's architectural metaphor onto digital infrastructures without attending to the qualitative differences that computational processing introduces. The capacity for algorithmic analysis of population-level data potentially enables forms of predictive governance that exceed the temporal constraints of traditional panoptic observation. Nevertheless, one might reasonably question whether such technological capabilities necessarily translate into effective governmental control, or whether they may paradoxically generate new forms of resistance and counter-surveillance.

Urban governance in the twenty-first century increasingly relies upon networked sensor technologies, facial recognition systems, and predictive analytics platforms that collectively constitute what could be characterised as an ambient surveillance infrastructure. Cities such as Singapore, Dubai, and increasingly London have invested substantially in these technological assemblages, positioning them as essential components of efficient urban management. Proponents argue that such systems enhance public safety, optimise resource allocation, and enable more responsive municipal services. Critics, however, suggest that these justifications may obscure more troubling implications for civil liberties and democratic accountability.

This research seeks to interrogate the discursive formations through which digital surveillance in urban contexts is legitimised, contested, and negotiated. Rather than adopting an a priori normative stance regarding the desirability or otherwise of such technologies, this study endeavours to map the complex terrain of competing rationalities that inform their deployment. It appears increasingly evident that simplistic binaries of freedom versus security inadequately capture the nuanced ways in which citizens navigate and make sense of their surveilled environments. The empirical component of this thesis therefore aims to foreground the lived experiences and interpretive frameworks of urban residents who find themselves enmeshed within these digital infrastructures.

The methodological approach adopted herein combines critical discourse analysis with ethnographic sensibilities, recognising that technologies are never merely technical artefacts but are always already embedded within social, political, and economic relations. While quantitative metrics of surveillance density or camera placement may provide useful contextual information, they arguably fail to illuminate the phenomenological dimensions of being watched. This thesis therefore privileges qualitative methods that enable deeper engagement with the meanings that subjects attach to their experiences of urban surveillance.

Several key research questions guide this investigation. Firstly, how do municipal authorities discursively frame digital surveillance as a governmental necessity, and what competing narratives challenge or complicate these framings? Secondly, in what ways do urban residents perceive, experience, and respond to the presence of digital surveillance technologies in their daily lives? Thirdly, what forms of resistance, adaptation, or accommodation emerge in response to the expanding surveillance apparatus? Finally, what implications might these findings have for theoretical understandings of power, governance, and citizenship in digitally mediated urban contexts?

The structure of this thesis proceeds as follows. Chapter Two provides a comprehensive review of relevant literature, situating the research within broader scholarly conversations regarding surveillance studies, urban governance, and critical data studies. Chapter Three elaborates the methodological framework, justifying the selection of research methods and discussing ethical considerations. Chapter Four presents the empirical findings, organised thematically around the key research questions. Chapter Five discusses these findings in relation to the theoretical framework, identifying areas of convergence and tension. Finally, Chapter Six offers conclusions and considers the implications for policy and future research.

It should perhaps be acknowledged at the outset that this research is conducted from a position that views uncritical technological solutionism with a degree of scepticism. This is not to suggest that digital technologies cannot serve beneficial purposes, but rather to insist upon the importance of subjecting their deployment to rigorous critical scrutiny. The normative commitment underlying this thesis is to the preservation and enhancement of democratic accountability, even as the meaning of such accountability may itself require renegotiation in light of technological change.

The significance of this research extends beyond purely academic considerations. As cities worldwide continue to invest in smart infrastructure and data-driven governance, questions regarding privacy, consent, and the distribution of power become matters of pressing public concern. This thesis aims to contribute to these vital conversations by providing empirically grounded insights into how digital surveillance is experienced and contested at the grassroots level. Ultimately, it is hoped that this research may inform more reflexive and accountable approaches to the integration of digital technologies within urban governance frameworks.

In conclusion, this introductory chapter has outlined the central problematic animating this thesis, introduced the key conceptual resources that will be deployed, and previewed the structure of the argument to follow. The subsequent chapters will develop these themes in greater detail, working towards a nuanced understanding of the digital panopticon as both a governmental technology and a site of ongoing social contestation.
"""

    for para in content.strip().split('\n\n'):
        p = doc.add_paragraph(para.strip())
        p.paragraph_format.first_line_indent = Pt(36)
        p.paragraph_format.space_after = Pt(12)

    doc.save(DRAFTS_DIR / "synthetic_intro.docx")
    word_count = len(content.split())
    print(f"Created synthetic_intro.docx ({word_count} words)")
    return word_count


def create_synthetic_lit_review():
    """
    Generate synthetic_lit_review.docx: 1,000 words reviewing Foucault and Zuboff
    with proper citations.
    """
    doc = Document()

    title = doc.add_heading("Chapter 2: Literature Review", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Theoretical Foundations: From Disciplinary Power to Surveillance Capitalism", level=2)

    content = """
The theoretical landscape informing contemporary surveillance studies is characterised by a productive tension between classical accounts of disciplinary power and emergent frameworks addressing the specificities of digital capitalism. This chapter critically examines the contributions of Michel Foucault and Shoshana Zuboff as foundational yet distinct theoretical resources for understanding the digital panopticon in urban governance contexts. While Foucault's (1977) analysis of disciplinary mechanisms provides essential conceptual tools, Zuboff's (2019) more recent intervention arguably captures dimensions of datafied power that exceed Foucauldian analytics.

Foucault's seminal work Discipline and Punish articulates a sophisticated account of how modern societies transitioned from spectacular punishment to normalising discipline. The architectural figure of the panopticon serves as both historical example and analytical metaphor, illustrating how power may operate through the internalisation of surveillance (Foucault, 1977). For Foucault, the genius of panoptic design lies in its capacity to induce "a state of conscious and permanent visibility that assures the automatic functioning of power" (Foucault, 1977, p. 201). Individuals subjected to panoptic observation become, in effect, their own overseers, adjusting their behaviour in accordance with anticipated scrutiny.

However, several scholars have questioned the direct applicability of Foucauldian analytics to contemporary digital surveillance regimes. Lyon (2006) suggests that while the panopticon metaphor usefully highlights the disciplinary dimensions of surveillance, it may obscure the ways in which digital technologies enable qualitatively different forms of monitoring. The computational capacity to aggregate, analyse, and act upon data at unprecedented scales introduces what Deleuze (1992) characterised as "societies of control," wherein regulation operates through continuous modulation rather than discrete disciplinary enclosures. Nevertheless, it would seem premature to entirely abandon Foucauldian insights, which retain considerable explanatory power when appropriately adapted.

Zuboff's (2019) The Age of Surveillance Capitalism provides perhaps the most comprehensive recent attempt to theorise the distinctive features of datafied power. Zuboff argues that major technology corporations have pioneered a novel accumulation logic that treats human experience as raw material for behavioural prediction products. This "surveillance capitalism" operates through asymmetric extraction of personal data, rendering individual subjects increasingly transparent to corporate actors while those actors themselves remain largely opaque (Zuboff, 2019). The implications for privacy and autonomy are, according to Zuboff, profoundly troubling.

Critically, Zuboff distinguishes surveillance capitalism from both traditional capitalism and state surveillance, positioning it as an unprecedented mutation in the history of capitalism itself. She contends that while states may participate in and benefit from surveillance infrastructures, the driving force behind their expansion is primarily economic rather than governmental (Zuboff, 2019). This framing has attracted both enthusiastic endorsement and critical scrutiny. Some scholars suggest that Zuboff underestimates the continuities between contemporary data practices and longer histories of capitalist surveillance (Morozov, 2019). Others argue that her focus on individual privacy may neglect collective and political dimensions of the problem (Cohen, 2019).

For the purposes of this thesis, both Foucauldian and Zuboffian frameworks offer valuable, if partial, perspectives. Foucault's (1977) attention to the productive dimensions of power—how surveillance not merely constrains but actively constitutes subjects—provides essential tools for understanding urban governance as a formative rather than merely repressive practice. Simultaneously, Zuboff's (2019) emphasis on the economic logics driving surveillance expansion helps illuminate why municipal authorities increasingly partner with private technology vendors and adopt data-intensive governance models.

A synthesis of these approaches might recognise that contemporary urban surveillance operates at the intersection of disciplinary and capitalist rationalities. Smart city initiatives, for instance, frequently combine governmental objectives such as public safety with commercial interests in data valorisation (Kitchin, 2014). The result is a hybrid regime in which citizens are simultaneously subjects of governmental discipline and sources of extractable data value. This dual positioning arguably complicates both Foucauldian accounts that foreground state power and Zuboffian narratives centred on corporate actors.

Furthermore, critical data studies scholars have emphasised the importance of attending to the materiality and infrastructure of surveillance systems (Parks and Starosielski, 2015). Abstract theorisations of power, whether Foucauldian or otherwise, may benefit from grounding in the concrete technical assemblages through which surveillance is enacted. Sensors, cables, servers, and algorithms are not merely neutral instruments but active participants in shaping how surveillance operates (Bowker and Star, 1999). This insight informs the methodological approach adopted in this thesis, which seeks to trace connections between theoretical frameworks and empirical observations of surveillance infrastructures.

The literature on resistance to surveillance provides another crucial dimension of the theoretical landscape. Drawing on Foucault's (1978) insistence that "where there is power, there is resistance," scholars have documented diverse tactics through which subjects contest, evade, or subvert surveillance regimes (Marx, 2003). From technical countermeasures such as encryption and anonymisation to performative strategies of sousveillance and obfuscation, the repertoire of resistance is varied and evolving (Brunton and Nissenbaum, 2015). This thesis takes seriously the agential capacities of urban residents, refusing to position them as passive objects of surveillance power.

In conclusion, this literature review has surveyed key theoretical resources for understanding digital surveillance in urban governance contexts. The works of Foucault (1977) and Zuboff (2019) provide complementary perspectives that, when synthesised, illuminate the complex interplay of disciplinary and economic rationalities shaping contemporary surveillance regimes. Subsequent chapters will draw upon these frameworks while remaining attentive to their limitations and the need for empirical grounding. The theoretical conversation reviewed here establishes the conceptual foundations upon which this thesis will build its original contribution to surveillance studies scholarship.
"""

    for para in content.strip().split('\n\n'):
        p = doc.add_paragraph(para.strip())
        p.paragraph_format.first_line_indent = Pt(36)
        p.paragraph_format.space_after = Pt(12)

    doc.save(DRAFTS_DIR / "synthetic_lit_review.docx")
    word_count = len(content.split())
    print(f"Created synthetic_lit_review.docx ({word_count} words)")
    return word_count


def create_synthetic_methodology():
    """
    Generate synthetic_methodology.docx: 800 words describing qualitative case study
    with Oxford Brookes Ethics Committee mentions.
    """
    doc = Document()

    title = doc.add_heading("Chapter 3: Methodology", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("A Qualitative Case Study Approach", level=2)

    content = """
This chapter elaborates the methodological framework guiding this investigation, justifying the selection of a qualitative case study approach and addressing the ethical considerations that inform the research design. The methodological choices outlined herein are grounded in the epistemological commitments articulated in the preceding chapters, recognising that knowledge production is always situated and partial rather than objective and complete.

The research adopts a qualitative case study methodology, which appears particularly well-suited to the exploratory and interpretive aims of this thesis. Case study research enables in-depth examination of complex social phenomena within their real-world contexts, attending to the particularities that may be obscured by more abstract methodological approaches (Yin, 2018). Given that this thesis seeks to understand how digital surveillance is experienced and negotiated within specific urban environments, the case study method offers appropriate tools for capturing contextual richness and nuance.

The selection of London as the primary case study site reflects both pragmatic and substantive considerations. London arguably represents one of the most intensively surveilled cities in the democratic world, with extensive networks of CCTV cameras, facial recognition deployments, and smart city initiatives. Furthermore, the city has been the site of significant public debate regarding surveillance, providing rich discursive material for analysis. Access considerations also informed this selection, as the researcher's location in the United Kingdom facilitated sustained engagement with the research site.

Data collection proceeded through multiple methods, triangulating sources to enhance the credibility and depth of findings. Semi-structured interviews constituted the primary data source, conducted with a purposive sample of urban residents, civil society activists, municipal officials, and technology professionals. Interview protocols were designed to elicit participants' experiences, perceptions, and interpretations of urban surveillance, while remaining sufficiently flexible to follow emergent themes. A total of twenty-five interviews were conducted over an eight-month period, with sessions ranging from forty-five minutes to two hours in duration.

Complementing the interview data, documentary analysis examined relevant policy documents, media coverage, and promotional materials associated with London's surveillance infrastructure. This analysis provided contextual understanding of the official discourses framing digital surveillance and enabled identification of tensions between policy rhetoric and lived experience. Additionally, observational fieldwork was conducted in public spaces characterised by visible surveillance presence, generating field notes that informed interpretation of interview data.

Ethical considerations were paramount throughout the research process. Full ethical approval was obtained from the Oxford Brookes Ethics Committee prior to commencement of data collection, ensuring that the research design met established standards for the protection of human participants. Informed consent was secured from all interview participants, who were provided with detailed information sheets explaining the purposes of the research, the voluntary nature of participation, and their right to withdraw at any time without consequence.

Given the potentially sensitive nature of discussing surveillance experiences, particular attention was paid to confidentiality and anonymisation. Participant identities have been protected through the use of pseudonyms, and identifying details have been altered or omitted where necessary. Interview recordings were stored securely using encryption, and access was restricted to the research team. The Oxford Brookes Ethics Committee specifically reviewed the data management protocols to ensure compliance with GDPR requirements and institutional guidelines.

The analytical approach drew upon thematic analysis as elaborated by Braun and Clarke (2006), involving systematic coding of interview transcripts and documentary materials. Initial open coding identified recurring patterns and concepts, which were subsequently organised into broader thematic categories. Throughout this process, analytical memos were maintained to document interpretive decisions and emerging insights. The analysis proceeded iteratively, with ongoing dialogue between data and theory informing the development of the final thematic structure.

Reflexivity constitutes an essential component of qualitative research practice, acknowledging that the researcher is not a neutral observer but an active participant in knowledge construction (Finlay, 2002). As a researcher positioned within a particular social location—a doctoral student at a British university with prior professional experience in the technology sector—certain perspectives and blind spots inevitably shape the interpretation of data. Efforts to mitigate researcher bias included regular supervision discussions, peer debriefing, and the maintenance of a reflexive journal throughout the fieldwork period.

Several limitations of the methodology should be acknowledged. The case study design, while enabling depth of analysis, necessarily constrains generalisability to other urban contexts. Findings from the London case cannot be straightforwardly extrapolated to cities with different surveillance configurations or cultural contexts. Furthermore, the purposive sampling strategy, while appropriate for exploratory research, may have introduced selection biases that affect the representativeness of participant perspectives. These limitations are acknowledged as inherent features of qualitative inquiry rather than methodological failures.

In summary, this chapter has outlined a qualitative case study methodology grounded in interpretivist epistemology and attentive to ethical responsibilities. The subsequent chapters will present and discuss the findings generated through this methodological framework, contributing to scholarly understanding of digital surveillance in urban governance.
"""

    for para in content.strip().split('\n\n'):
        p = doc.add_paragraph(para.strip())
        p.paragraph_format.first_line_indent = Pt(36)
        p.paragraph_format.space_after = Pt(12)

    doc.save(DRAFTS_DIR / "synthetic_methodology.docx")
    word_count = len(content.split())
    print(f"Created synthetic_methodology.docx ({word_count} words)")
    return word_count


def create_mock_results():
    """
    Generate mock_results.csv with 20 rows of interview codes.
    """
    themes = [
        "Ambient Awareness",
        "Normalisation of Surveillance",
        "Privacy Trade-offs",
        "Safety vs Freedom",
        "Algorithmic Opacity",
        "Data Extraction Concerns",
        "Resistance Tactics",
        "Technological Fatalism",
        "Institutional Trust",
        "Spatial Avoidance"
    ]

    quotes = [
        "You just stop thinking about the cameras after a while, they become part of the background.",
        "I know they're watching but honestly if it keeps my kids safe, I can live with it.",
        "The thing that bothers me is not knowing what they do with all this information.",
        "Sometimes I deliberately take longer routes to avoid the main surveillance areas.",
        "They say it's for our protection but who protects us from them?",
        "I've started wearing different clothes and hats when I go certain places.",
        "My generation grew up with this, we don't really know anything different.",
        "The council keeps saying it's all anonymised but I don't believe them.",
        "I feel like we've already lost the battle, there's cameras everywhere now.",
        "It's not just CCTV anymore, it's your phone, your oyster card, everything.",
        "I trust the government more than the private companies to be honest.",
        "When I see those cameras I do wonder who's actually looking at the footage.",
        "People change how they behave in public now, everyone's more self-conscious.",
        "The facial recognition stuff really worries me, that's a step too far.",
        "Some areas of London feel completely surveilled, you can't escape it.",
        "I think most people have just accepted this as the price of living in a city.",
        "There's no real consent, you walk outside and you're automatically in the system.",
        "They target certain communities more than others, it's obvious if you pay attention.",
        "I've thought about moving somewhere with less surveillance but everywhere's the same now.",
        "At least with human officers you could reason with them, algorithms don't listen."
    ]

    rows = []
    for i in range(20):
        rows.append({
            "Participant_ID": f"P{i+1:02d}",
            "Theme": random.choice(themes),
            "Quote_Snippet": quotes[i],
            "Interview_Date": f"2025-{random.randint(3,8):02d}-{random.randint(1,28):02d}",
            "Location": random.choice(["Central London", "East London", "South London", "North London", "West London"]),
            "Age_Range": random.choice(["18-25", "26-35", "36-45", "46-55", "56+"]),
            "Sentiment": random.choice(["Negative", "Ambivalent", "Resigned", "Concerned", "Neutral"])
        })

    csv_path = DATA_DIR / "mock_results.csv"
    with open(csv_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=rows[0].keys())
        writer.writeheader()
        writer.writerows(rows)

    print(f"Created mock_results.csv (20 rows)")
    return len(rows)


def create_supervisor_feedback():
    """
    Generate supervisor_email.txt with critique about methodology.
    """
    feedback = """From: Dr. Sarah Mitchell <s.mitchell@brookes.ac.uk>
To: PhD Candidate
Subject: Feedback on Chapters 1-3 Draft
Date: 15 January 2026

Dear Candidate,

Thank you for submitting your draft chapters. I have reviewed them carefully and have the following observations.

CRITICAL FEEDBACK:

Your methodology chapter is currently too descriptive. You need to engage more substantively with the power dynamics mentioned in the Literature Review. Specifically:

1. The discussion of Foucault's panopticon in Chapter 2 raises important questions about power asymmetries, but these are not adequately addressed in your methodological approach. How does your interview methodology account for the power differential between researcher and researched? You mention reflexivity briefly but this needs much more development.

2. The connection between Zuboff's surveillance capitalism framework and your data collection methods is unclear. If you're arguing that datafied power operates through extraction, shouldn't your methodology explicitly address how you avoid replicating extractive logics in your own research practice?

3. The case study justification is somewhat thin. London is described as "one of the most intensively surveilled cities" but this claim needs evidential support. What metrics are you using? How does London compare to other possible case sites?

POSITIVE OBSERVATIONS:

- Your theoretical synthesis of Foucault and Zuboff in Chapter 2 is sophisticated and original. This is a strength of the thesis.
- The hedging language and academic tone are appropriate for doctoral work.
- The structure is clear and the argument flows logically between sections.

MINOR CORRECTIONS:

- Check citation format on p. 12 (Foucault, 1977) - should include page number for direct quote.
- Some sentences in the Introduction exceed 50 words - consider breaking these up.
- Consistency in terminology: you alternate between "surveillance" and "monitoring" - pick one and use consistently.

Please address the critical feedback as a priority before our next supervisory meeting.

Best regards,
Dr. Mitchell
Senior Lecturer in Sociology
Oxford Brookes University
"""

    feedback_path = FEEDBACK_DIR / "supervisor_email.txt"
    with open(feedback_path, 'w', encoding='utf-8') as f:
        f.write(feedback)

    print(f"Created supervisor_email.txt")
    return 1


def run_dna_extraction():
    """Run the DNA engine on the synthetic files."""
    import subprocess
    import sys

    print("\nRunning DNA extraction...")
    result = subprocess.run(
        [sys.executable, str(ROOT_DIR / "core" / "dna_engine.py")],
        capture_output=True,
        text=True,
        cwd=str(ROOT_DIR)
    )

    if result.returncode == 0:
        print("DNA extraction completed successfully")
        print(result.stdout[-500:] if len(result.stdout) > 500 else result.stdout)
    else:
        print(f"DNA extraction had issues: {result.stderr}")

    return result.returncode == 0


def main():
    """Generate all synthetic research environment files."""
    print("=" * 60)
    print("PHDx Synthetic Research Environment Generator")
    print("=" * 60)
    print()

    total_words = 0

    # Create drafts
    print("Creating synthetic draft documents...")
    total_words += create_synthetic_intro()
    total_words += create_synthetic_lit_review()
    total_words += create_synthetic_methodology()

    print()

    # Create mock data
    print("Creating mock interview data...")
    create_mock_results()

    print()

    # Create feedback
    print("Creating supervisor feedback...")
    create_supervisor_feedback()

    print()

    # Run DNA extraction
    dna_success = run_dna_extraction()

    print()
    print("=" * 60)

    # Verify
    files_created = [
        DRAFTS_DIR / "synthetic_intro.docx",
        DRAFTS_DIR / "synthetic_lit_review.docx",
        DRAFTS_DIR / "synthetic_methodology.docx",
        DATA_DIR / "mock_results.csv",
        FEEDBACK_DIR / "supervisor_email.txt",
        DATA_DIR / "author_dna.json"
    ]

    all_exist = all(f.exists() for f in files_created)

    if all_exist and dna_success:
        print()
        print("=" * 60)
        print("        READY FOR TEST")
        print("=" * 60)
        print()
        print(f"Total synthetic words created: {total_words:,}")
        print(f"Files generated: {len(files_created)}")
        print()
        print("Synthetic environment includes:")
        print("  - 3 DOCX drafts (Intro, Lit Review, Methodology)")
        print("  - 20 rows of mock interview data")
        print("  - Supervisor feedback with critique")
        print("  - Author DNA profile extracted")
        print()
        print("You can now test all PHDx features!")
        print("=" * 60)
    else:
        print("Some files may not have been created. Please check errors above.")
        for f in files_created:
            status = "OK" if f.exists() else "MISSING"
            print(f"  [{status}] {f.name}")


if __name__ == "__main__":
    main()
